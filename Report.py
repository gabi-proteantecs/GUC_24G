import copy
import os
import time

import matplotlib.pyplot as plt  # Importing the matplotlib.pyplot
import numpy
import numpy as np
import openpyxl
import pandas as pd
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT  # 导入表格对齐方式
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor
from openpyxl.styles import Alignment, Border, Font, Side


class Graph:
    def __init__(self, gui):
        self.gui = gui
        self.i2c = None

    def Word_start_head(self):
        doc = Document()
        # print(TestDataRate_val)

        doc.styles["Normal"].font.name = "Arial"
        doc.styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
        doc.styles["Normal"].font.size = Pt(10)
        doc.styles["Normal"].font.color.rgb = RGBColor(0, 0, 0)
        doc.add_paragraph().add_run("X.X.X.X   Test Result").bold = True
        doc.add_paragraph("Test result are show in below.")
        doc.add_paragraph().add_run(
            "Table X.X.X : " + Intreface_name + TestFunction_val + "\n\n"
        ).bold = True
        doc.save(r"Test Report.docx")

    def Word_head(self):
        time.sleep(0.2)
        doc = Document(r"Test Report.docx")
        # time.sleep(0.1)
        doc.styles["Normal"].font.name = "Arial"
        doc.styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
        doc.styles["Normal"].font.size = Pt(10)
        doc.styles["Normal"].font.color.rgb = RGBColor(0, 0, 0)
        # doc.add_paragraph().add_run('X.X.X.X   Test Result').bold = True
        doc.add_paragraph("Test result are show in below.")
        doc.add_paragraph().add_run(
            "Table X.X.X : " + Intreface_name + TestFunction_val + "\n\n"
        ).bold = True
        doc.save(r"Test Report.docx")

    def Word_table(self, num):
        # Edit word report table
        time.sleep(0.2)
        doc = Document(r"Test Report.docx")
        # time.sleep(0.1)

        Func_Num = num
        Txt_Data = open(r"Test_Report_Log.txt")
        text = []
        for line in Txt_Data:
            text.append(line)
        # print(text)
        for i in range(999):  # check test item number
            try:
                lens = round(len(str(text[i])))
                # print(lens)
            except Exception:
                break
        global TestItem_number

        # Have "to" in TestSpec_val?
        # print(TestSpec_val)
        row_num = (
            TestItem_number + 4
        )  # Test item + table title(ex:NV - NT) + Spec. Max + Spec. Min
        # print(TestItem_number)
        cols_num = 7

        # edit test condition
        table = doc.add_table(
            rows=row_num, cols=cols_num, style="Table Grid"
        )  # 创建带边框的表格
        table.cell(0, 0).paragraphs[0].add_run("").bold = True
        table.cell(0, 1).paragraphs[0].add_run(
            "LV - LT"
        ).bold = True  #  table.cell(row, cols)
        table.cell(0, 2).paragraphs[0].add_run("LV - HT").bold = True
        table.cell(0, 3).paragraphs[0].add_run("NV - NT").bold = True
        table.cell(0, 4).paragraphs[0].add_run("HV - LT").bold = True
        table.cell(0, 5).paragraphs[0].add_run("HV - HT").bold = True
        table.cell(0, 6).paragraphs[0].add_run("Unit").bold = True

        # edit test chip version (ex:TT01) table
        # print(TestSpec_val)
        # print(row_num)
        if "to" in TestSpec_val:  # Have "to" in TestSpec_val?
            for i in range(TestItem_number):
                # print(Sweep_Array_1Darray[i])
                table.cell(i + 1, 0).paragraphs[0].add_run(
                    Sweep_Array_1Darray[i]
                ).bold = True  # test chip corsor/number table title
                table.cell(i + 1, 6).paragraphs[0].add_run(
                    TestUnit_val
                ).bold = True  # test chip 'Unit' table title #  table.cell(row, cols)
            table.cell(i + 2, 0).paragraphs[0].add_run(
                "Spec. Min"
            ).bold = True  # Spec. Max #  table.cell(row, cols)
            table.cell(i + 2, 6).paragraphs[0].add_run(
                TestUnit_val
            ).bold = True  # test chip 'Unit' table title #  table.cell(row, cols)
            table.cell(i + 3, 0).paragraphs[0].add_run(
                "Spec. Max"
            ).bold = True  # Spec. Min #  table.cell(row, cols)
            table.cell(i + 3, 6).paragraphs[0].add_run(
                TestUnit_val
            ).bold = True  # test chip 'Unit' table title #  table.cell(row, cols)
            for col in range(5):
                SpecRange = TestSpec_val.split("to")
                table.cell(TestItem_number + 1, col + 1).paragraphs[0].add_run(
                    SpecRange[0]
                ).bold = False
                SpecMax = SpecRange[1].strip("\n")
                # print(SpecMax)
                table.cell(TestItem_number + 2, col + 1).paragraphs[0].add_run(
                    SpecMax
                ).bold = False
            for i in range(TestItem_number + 1):  # test item num + spec. Max
                table.rows[i + 1].height = Pt(25)
        elif "NA" in TestSpec_val:
            for i in range(TestItem_number):
                # print(TestItem_number)
                table.cell(i + 1, 0).paragraphs[0].add_run(
                    Sweep_Array_1Darray[i]
                ).bold = True  # test chip corsor/number table title
                table.cell(i + 1, 6).paragraphs[0].add_run(
                    TestUnit_val
                ).bold = True  # test chip 'Unit' table title #  table.cell(row, cols)
        else:
            for i in range(TestItem_number):
                # print(TestItem_number)
                table.cell(i + 1, 0).paragraphs[0].add_run(
                    Sweep_Array_1Darray[i]
                ).bold = True  # test chip corsor/number table title
                table.cell(i + 1, 6).paragraphs[0].add_run(
                    TestUnit_val
                ).bold = True  # test chip 'Unit' table title #  table.cell(row, cols)
            table.cell(i + 2, 0).paragraphs[0].add_run(
                "Spec. Max"
            ).bold = True  # Spec. Max #  table.cell(row, cols)

            table.cell(i + 2, 6).paragraphs[0].add_run(
                TestUnit_val
            ).bold = True  # test chip 'Unit' table title #  table.cell(row, cols)
            for col in range(5):
                SpecMax = TestSpec_val.strip("\n")
                table.cell(TestItem_number + 1, col + 1).paragraphs[0].add_run(
                    SpecMax
                ).bold = False
            for i in range(TestItem_number + 1):  # test item num + spec. Max
                table.rows[i + 1].height = Pt(25)

        # edit test value table
        for col in range(5):  # 5 : LV - LT , HV - LT , NV - NT , LV - HT , HV - HT
            condition_array = ["LV - LT", "LV - HT", "NV - NT", "HV - LT", "HV - HT"]
            condition_val = condition_array[col]
            # print(condition_val)
            for row in range(TestItem_number):
                # print(TestItem_number)
                Sweep_val = Sweep_Array_1Darray[row]
                # print(condition_val,Sweep_val)
                for i in range(999):
                    try:
                        # print(text[i])
                        # print(i)
                        s = text[i + 1]
                        DataRate_val = float(s.split(",")[5])
                        # print(DataRate_val)
                        if (
                            s.find(condition_val) == -1
                        ):  # Check LV - LT , HV - LT , NV - NT , LV - HT , HV - HT
                            pass
                        else:
                            # print('value',s)
                            if s.find(Sweep_val) == -1:  # check test item number
                                my_array_val = "NA"
                                table.cell(row + 1, col + 1).paragraphs[0].add_run("")
                            else:
                                if DataRate_val != TestDataRate_val:
                                    pass
                                else:
                                    my_array = s.split(",")
                                    my_array_val = my_array[Func_Num]
                                    # print(my_array_val)
                                    table.cell(row + 1, col + 1).paragraphs[0].add_run(
                                        my_array_val
                                    )  # table.cell(row, cols) Start write test value
                                    doc.save(r"Test Report.docx")
                    except Exception:
                        break

        # read table and Fill up blank -->NA
        for col in range(5):
            for row in range(TestItem_number):
                tables = doc.tables
                result = table.cell(row + 1, col + 1).text
                # print(result)
                if result == "":
                    # print(col)
                    table.cell(row + 1, col + 1).paragraphs[0].add_run("NA")
                else:
                    pass

        Item_Num = i
        for r in range(row_num):  # 循环将每一行，每一列都设置为居中
            for c in range(cols_num):
                table.cell(r, c).paragraphs[
                    0
                ].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
                table.cell(r, c).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                table.cell(r, c).width = Pt(9999)
                table.rows[r].height = Pt(15)

        table.rows[0].height = Pt(40)
        doc.save(r"Test Report.docx")

    def Word_picture(self):
        time.sleep(0.2)
        doc = Document(r"Test Report.docx")
        # time.sleep(0.1)
        section = doc.sections[0]
        section.left_margin = Cm(1.27)
        section.right_margin = Cm(1.27)
        section.top_margin = Cm(1.27)
        section.bottom_margin = Cm(1.27)
        string = "文字内容"
        images = "Graph.png"  # 保存在本地的图片
        doc.add_paragraph("\n")  # 添加文字
        doc.add_picture("Graph.png", width=Cm(18), height=Cm(8))
        last_paragraph = doc.paragraphs[
            -1
        ]  # 段落屬性，在這裏代表每一行，一共三行，-1為最後一行
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.styles["Normal"].font.name = "Arial"
        doc.styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
        doc.styles["Normal"].font.size = Pt(10)
        doc.styles["Normal"].font.color.rgb = RGBColor(0, 0, 0)
        paragraph = doc.add_paragraph()
        paragraph_format = paragraph.paragraph_format
        paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.add_run(
            "Table X.X.X : " + Intreface_name + TestFunction_val + "\n\n"
        ).bold = True

        doc.add_page_break()

        doc.save(r"Test Report.docx")

    def Word_new_line(self, line_num):
        time.sleep(0.2)
        doc = Document(r"Test Report.docx")
        # time.sleep(0.1)
        for i in range(line_num):
            # print(line_num)
            doc.add_paragraph("")
        doc.save(r"Test Report.docx")

    def Word_DeleteRow(self):
        for s in range(300):
            doc = Document(r"Test Report.docx")
            try:
                table = doc.tables[s]
                for a in range(10):
                    for i in range(99):
                        try:
                            my_list = table.cell(i, 1).text
                            if my_list == "":
                                row2 = table.rows[i]
                                row2._element.getparent().remove(row2._element)
                                doc.save(r"Test Report.docx")
                            else:
                                pass
                        except Exception:
                            break
            except Exception:
                break

    def Word_next_page(self):
        time.sleep(0.2)
        doc = Document(r"Test Report.docx")
        # time.sleep(0.1)
        doc.add_page_break()

    def xls_save(self, Excel_Path, **kargs):
        excel_path = kargs.get("path", "A.xlsx")
        # # excel_sheet = kargs.get('sheet', 0)
        #
        # Excel_Path = load_workbook(filename=excel_path)
        # Select_sheet = Excel_Path[excel_sheet]
        # Excel_Path.save(excel_path)
        try:
            Excel_Path.save(excel_path)
        except Exception as e:
            if "Permission" in str(e):
                new_p = excel_path.rstrip(".xlsx") + "_bk1.xlsx"
                print(f"\033Cannot save {excel_path}, save {new_p} instead")
                Excel_Path.save(new_p)
                raise BaseException("Error")
            else:
                print(f"\033Save file error")
                print(e)
                raise BaseException("Error")
        # Excel_Path.close()

    def xls_close(self, **kargs):
        import os

        command = "taskkill /f /t /im EXCEL.exe"
        os.system(command)

    def xls_open(self, **kargs):
        sheet_name = kargs.get("xls_sheet", 0)
        from win32com.client import Dispatch

        xl = Dispatch("Excel.Application")
        xl.Visible = True  # otherwise excel is hidden

        # newest excel does not accept forward slash in path
        wb = xl.Workbooks.Open("\Test Report(EY0012A).xlsx")
        wb.Close()
        xl.Quit()

    def xls_check_status(self):
        pass

    def xls_ipxact(self, **kargs):
        Flow_S = kargs.get("Flow_S", "3.6.1")
        def_name = kargs.get("def_name", "")
        defFile_en = kargs.get("def_name", 1)
        reg_source = kargs.get("reg_source", "xlsx")  # xls or txt
        col_val = kargs.get("col_val", 5)  # 5 is Def col number in IPXACT sheet
        col_fow = kargs.get("col_fow", 500)

        if reg_source == "xlsx":
            if defFile_en:
                path = "TestTools/IPXACT_def/" + def_name + ".txt"
                f = open(path, mode="w")
                f.write("")
                f.close()

            wb = openpyxl.load_workbook("Test Report\Test Report.xlsx")
            sh = wb["IPXACT"]

            # check excel register row number
            Num = []
            Col_s = 1
            for Reg in [Flow_S]:
                for row1 in range(9999):
                    try:
                        xls_row = row1
                        val = sh.cell(row=xls_row, column=Col_s).value
                        if val.find(Reg) != -1:
                            break
                        else:
                            pass
                    except Exception as e:
                        val = "Blank"
                        pass
                    if row1 == 999:
                        val = f"Error : Please check IPXACT flow"
                    # print(val)
                Num += [xls_row + 1]
            row_S = Num[0] + 1

            Num = []
            for Reg in ["Done"]:
                for row1 in range(9999):
                    try:
                        xls_row = row1 + row_S
                        val = sh.cell(row=xls_row, column=Col_s).value
                        if val.find(Reg) != -1:
                            break
                        else:
                            pass
                    except Exception as e:
                        val = "Blank"
                        pass
                    if row1 == 999:
                        val = f"Error : Please check IPXACT flow"
                    # print(val)
                Num += [xls_row]
            row_E = Num[0]

            reg_arr = []
            reg = ""
            # read register offset, S_Bit, B_Leng, Value
            Ref_rgister_col = 0  # excel register value col number in ipxact sheet
            for i in range(row_E - row_S):
                offset = sh.cell(row=row_S + i, column=Col_s + 1).value
                S_Bit = sh.cell(row=row_S + i, column=Col_s + 2).value
                B_Leng = sh.cell(row=row_S + i, column=Col_s + 3).value

                Value = sh.cell(row=row_S + i, column=col_val).value
                slice_flow = sh.cell(row=row_S + i, column=col_fow).value
                if (
                    offset == "skip"
                    or offset == "Offset_Address"
                    or (str(offset).strip()) == "None"
                ):
                    pass
                else:
                    reg = f"{offset}, {S_Bit}, {B_Leng}, {Value}, {slice_flow}"
                    reg_arr += [reg]

                    if defFile_en:
                        path = "TestTools/IPXACT_def/" + def_name + ".txt"
                        f = open(path, mode="a+")
                        f.write(reg + "\n")
                        f.close()

        # register source from def file
        path = "TestTools/IPXACT_def/" + def_name + ".txt"
        f = open(path)
        if reg_source == "txt":
            reg_arr = []
            for line in f.readlines():
                reg_arr.append(line)
        f.close

        # print(f'- {def_name}', flush=True)
        return reg_arr

    def txt_log_check(self, **kargs):
        txt_path = kargs.get("txt_path", "NA")
        rst_test = kargs.get("rst_test", 0)

        excel_path = self.gui.BIST_Excel_path.GetPath()
        if excel_path != "":
            aaa = excel_path.split("\\")
            a = len(aaa)
            delete_str = aaa[a - 1]
            txt_folder = excel_path.rstrip(delete_str)
            txt_path = txt_folder + txt_path
        else:
            pass

        # print('\n\n\n\n------------------------------------ BIST FAIL Check ------------------------------------')
        fail_list = pass_list = []
        pattern_last = ""
        vco_fail = 0
        err_chk_fail = 0
        with open(txt_path, "r") as f:
            Chip_Info_list = f.readlines()
            index = 0
            for log in Chip_Info_list:
                buffer1 = log.find("BIST FAIL")
                buffer2 = log.find("BIST PASS")
                # info = (index, buffer1, log)  # index : txt row number
                buffer3 = log.find(" times fail in ")
                buffer4 = log.find("lol = 0x1")
                buffer5 = log.find("error check failed")
                if buffer4 != -1:
                    vco_fail = 1
                    for i in range(10):
                        chk_test = Chip_Info_list[index - i]
                        T_F = chk_test.find(" PLL_en")
                        if T_F != -1:
                            vco_info = (chk_test.split(" PLL_en"))[0]
                            vco_info = vco_info.replace("die", "_Die ")
                            break
                if buffer5 != -1:
                    err_chk_fail = 1
                    for i in range(40):
                        err_chk_info = Chip_Info_list[index - i]
                        T_F2 = err_chk_info.find("BIST check")
                        if T_F2 != -1:
                            err_chk_info = (err_chk_info.split("-BIST"))[0]
                            err_chk_info = (err_chk_info.split(" "))[-1]
                            break
                        else:
                            err_chk_info = "No_Find"
                            pass
                if buffer1 == -1 and buffer2 == -1:
                    pass
                else:
                    if buffer1 != -1:
                        slices = Chip_Info_list[index - 1]
                    else:
                        slices = ""
                    for i in range(40):
                        pattern = Chip_Info_list[index - i]
                        T_F = pattern.find("bist chk")
                        T_F2 = pattern.find("BIST check")
                        if T_F != -1:
                            # print(pattern)
                            pattern = (pattern.split(":"))[1]
                            pattern = (pattern.split(" bist"))[0]
                            pattern = (pattern.split(" "))[-1]
                            # print(pattern)
                            # print(index - i)
                            break
                        elif T_F2 != -1:
                            pattern = (pattern.split("-BIST"))[0]
                            pattern = (pattern.split(" "))[-1]
                            break
                        else:
                            pattern = "No_Find"
                            pass
                    for i in range(40):
                        if index + i >= len(Chip_Info_list):
                            break
                        Lan_Fail = Chip_Info_list[index + i]
                        T_F3 = Lan_Fail.find("Total Lane Fail NUM")
                        if T_F3 != -1:
                            fail_lane_num = (Lan_Fail.split("Total Lane Fail NUM = "))[
                                1
                            ]
                            fail_lane_num = (fail_lane_num.split("\n"))[0]
                            fail_lane_num = f" FN={fail_lane_num}"
                            break
                        else:
                            fail_lane_num = " FN=0"
                    for i in range(800):
                        if index + i >= len(Chip_Info_list):
                            break
                        Lan_Fail = Chip_Info_list[index + i]
                        T_F4 = Lan_Fail.find(" times fail in ")
                        if T_F4 != -1:
                            rst_fail_num = (Lan_Fail.split(" times fail in "))[0]
                            rst_fail_num = (rst_fail_num.split("get "))[1]
                            rst_fail_num = f" RST_FN={rst_fail_num}"
                            break
                        else:
                            fail_lane_num = " FN=0"
                            rst_fail_num = " RST_FN=0"
                    log = (log.split(":"))[0]
                    if slices == "":
                        if err_chk_fail == 1:
                            fail_str = (
                                err_chk_info.strip() + "_" + log.strip() + "_ERRFAIL"
                            )
                        elif vco_fail == 0:
                            fail_str = pattern.strip() + "_PASS"
                        elif vco_fail == 1:
                            fail_str = pattern.strip() + vco_info.strip() + "_LOL"
                        else:
                            fail_str = pattern.strip() + vco_info.strip() + "_LOL"
                    else:
                        fail_str = (
                            pattern.strip()
                            + "_"
                            + log.strip()
                            + "_"
                            + slices.strip()
                            + fail_lane_num
                            + rst_fail_num
                        )
                    # print(fail_str)
                    fail_list += [fail_str]
                index += 1

            # print(fail_list)
            fail_list = numpy.unique(fail_list).tolist()
        # print(fail_list)
        # self.PMAD_BIST_DATA_SEL = ['PRBS7', 'PRBS31', 'CLOCK', 'All-0', 'PRBS5', 'PRBS9', 'USER0_b15b0', 'RSV',
        #                   'N_PRBS7', 'N_PRBS31', 'N_CLOCK', 'All-1', 'N_PRBS5', 'N_PRBS9', 'N_USER0_b15b0', 'RSV']
        # self.PCS_BIST_DATA_SEL = ['PRBS32', 'USER_P0P0', 'PRBS32', 'USER_P0P1']
        chk_list = [
            "PRBS31",
            "PRBS7",
            "CLOCK",
            "All-0",
            "All-1",
            "USER0_b15b0",
            "PRBS32",
            "P0P1",
        ]
        chk_list = [
            "PRBS9",
            "PRBS7",
            "CLOCK",
            "All-0",
            "PRBS31",
            "USER0_b15b0",
            "PRBS32",
            "P0P1",
        ]
        F_PAT_list = []
        BIST_FN = []
        BIST_RstFN = []
        for ckp in chk_list:
            p_tested = -1
            F_slices_ck = []
            F_die_ck = []
            Fail_num_list = []
            Fail_rst_num = 0
            for i in range(len(fail_list)):
                T_F = fail_list[i].find(ckp)
                if T_F != -1:
                    if fail_list[i].find("PASS") == -1:
                        F_die = int(fail_list[i][(fail_list[i].find("Die") + 4)])
                        if (
                            fail_list[i].find("LOL") == -1
                            and fail_list[i].find("ERRFAIL") == -1
                        ):
                            sice_idx_s = fail_list[i].find("[") + 1
                            sice_idx_e = fail_list[i].find("]")
                            F_slices_info = fail_list[i][sice_idx_s:sice_idx_e].split(
                                ","
                            )
                            Fail_num_info = fail_list[i].split("FN=")[1]
                            Fail_num_info = int(Fail_num_info.split(" RST")[0])
                            Fail_rst_num = int(fail_list[i].split("RST_FN=")[1])
                            # # Fail_num_info = [x[0:2].rstrip('\n') for x in Fail_num_info_o[1:] ]
                            # Fail_num_info = [0,0,0,0,0,0,0,0]
                        if F_die not in F_die_ck:
                            F_die_ck.append(F_die)
                            F_slices = []
                            F_slices_ck.append(F_slices)
                            if (
                                fail_list[i].find("LOL") == -1
                                and fail_list[i].find("ERRFAIL") == -1
                            ):
                                Fail_num_list.append(Fail_num_info)
                                # Fail_num_list_ck.append(Fail_num_info)
                        else:
                            if (
                                fail_list[i].find("LOL") == -1
                                and fail_list[i].find("ERRFAIL") == -1
                            ):
                                F_slices = F_slices_ck[F_die_ck.index(F_die)]
                                if Fail_num_info > Fail_num_list[F_die_ck.index(F_die)]:
                                    Fail_num_list[F_die_ck.index(F_die)] = Fail_num_info
                                # Fail_num_list = Fail_num_list_ck[F_die_ck.index(F_die)]
                        if (
                            fail_list[i].find("LOL") == -1
                            and fail_list[i].find("ERRFAIL") == -1
                        ):
                            for s in F_slices_info:
                                if int(s) not in F_slices:
                                    F_slices.append(int(s))
                                    # Fail_num_list.append(int(Fail_num_info[F_slices_info.index(s)]))
                                # F_slices.append(int(s))
                    else:
                        pass
                    p_tested = 1
                else:
                    pass
            F_PAT = []
            for i in range(len(F_die_ck)):
                F_PAT += [f"D{F_die_ck[i]} {F_slices_ck[i]}"]
                # print(F_PAT)
            F_PAT_str = "\n".join(F_PAT)
            if p_tested == -1:
                F_PAT_list.append(-1)
                BIST_FN.append(0)
                BIST_RstFN.append(0)
            else:
                F_PAT_list.append(F_PAT_str)
                BIST_FN.append(sum(Fail_num_list))
                BIST_RstFN.append(Fail_rst_num)

        BIST_Slice = []
        BIST_val = []
        BIST_result = "Not Test"
        for pat_tested in F_PAT_list:
            if pat_tested == -1:
                BIST_Slice.append("")
                BIST_val.append("Not Test")
            else:
                BIST_Slice.append(pat_tested)
                if pat_tested == "":
                    BIST_val.append("PASS")
                    if BIST_result == "Not Test":
                        BIST_result = "PASS"
                else:
                    BIST_val.append("FAIL")
                    BIST_result = "FAIL"

        # BIST_PRBS31_Slice = BIST_Slice[0]
        # BIST_PRBS7_Slice  = BIST_Slice[1]
        # BIST_CLOCK_Slice  = BIST_Slice[2]
        # BIST_ALL0_Slice   = BIST_Slice[3]
        # BIST_ALL1_Slice   = BIST_Slice[4]
        # BIST_USR0_Slice   = BIST_Slice[5]
        # BIST_PRBS32_Slice = BIST_Slice[6]
        # BIST_POP1_Slice   = BIST_Slice[7]
        #
        # BIST_PRBS31_val = BIST_val[0]
        # BIST_PRBS7_val  = BIST_val[1]
        # BIST_CLOCK_val  = BIST_val[2]
        # BIST_ALL0_val   = BIST_val[3]
        # BIST_ALL1_val   = BIST_val[4]
        # BIST_USR0_val   = BIST_val[5]
        # BIST_PRBS32_val = BIST_val[6]
        # BIST_POP1_val   = BIST_val[7]
        # (F'{BIST_result};{BIST_PRBS31_val};{BIST_PRBS7_val};{BIST_CLOCK_val};{BIST_PRBS32_val};{BIST_POP1_val};'
        #  F'{BIST_PRBS31_Slice};{BIST_PRBS7_Slice};{BIST_CLOCK_Slice};{BIST_PRBS32_Slice};{BIST_POP1_Slice}')
        return_text_0 = f"{BIST_result};"
        return_text_1 = ""
        return_text_2 = ""
        return_text_3 = ""
        for i in range(len(BIST_Slice)):
            return_text_1 += f"{BIST_val[i]};"
            return_text_2 += f"{BIST_Slice[i]};"
            if rst_test == 0:
                return_text_3 += f"{BIST_FN[i]};"
            else:
                return_text_3 += f"{BIST_RstFN[i]};"

        """
        PRBS32_val = PRBS31_val = PRBS7_val = CLOCK_val = ''
        characters = "'Fail Slice = "
        with open(txt_path, "r") as f:
            buffer = ((f.read()).find('PRBS32'))
            if buffer != -1:
                if F_PRBS32_str == '':
                    BIST_PRBS32_val = "PASS"
                    BIST_PRBS32_Slice = ''
                else:
                    BIST_PRBS32_val = "FAIL"
                    F_PRBS32_str = ''.join(x for x in F_PRBS32_str if x not in characters)
                    BIST_PRBS32_Slice = F_PRBS32_str
            else:
                BIST_PRBS32_val = "Not Test"
                BIST_PRBS32_Slice = ''

        with open(txt_path, "r") as f:
            buffer = ((f.read()).find('PRBS31'))
            if buffer != -1:
                if F_PRBS31_str == '':
                    BIST_PRBS31_val = "PASS"
                    BIST_PRBS31_Slice = ''
                else:
                    BIST_PRBS31_val = "FAIL"
                    F_PRBS31_str = ''.join(x for x in F_PRBS31_str if x not in characters)
                    BIST_PRBS31_Slice = F_PRBS31_str
            else:
                BIST_PRBS31_val = "Not Test"
                BIST_PRBS31_Slice = ''

        with open(txt_path, "r") as f:
            buffer = ((f.read()).find('PRBS7'))
            if buffer != -1:
                if F_PRBS7_str == '':
                    BIST_PRBS7_val = "PASS"
                    BIST_PRBS7_Slice = ''
                else:
                    BIST_PRBS7_val = "FAIL"
                    F_PRBS7_str = ''.join(x for x in F_PRBS7_str if x not in characters)
                    BIST_PRBS7_Slice = F_PRBS7_str
            else:
                BIST_PRBS7_val = "Not Test"
                BIST_PRBS7_Slice = ''

        with open(txt_path, "r") as f:
            buffer = ((f.read()).find('CLOCK'))
            if buffer != -1:
                if F_CLOCK_str == '':
                    BIST_CLOCK_val = "PASS"
                    BIST_CLOCK_Slice = ''
                else:
                    BIST_CLOCK_val = "FAIL"
                    F_CLOCK_str = ''.join(x for x in F_CLOCK_str if x not in characters)
                    BIST_CLOCK_Slice = F_CLOCK_str
            else:
                BIST_CLOCK_val = "Not Test"
                BIST_CLOCK_Slice = ''
        """

        return return_text_0 + return_text_1 + return_text_3 + return_text_2[:-1]
        # return return_text_0 + return_text_1 + return_text_2[:-1]

    def txt_log_count_check(self, **kargs):
        txt_path = kargs.get("txt_path", "NA")

        excel_path = self.gui.BIST_Excel_path.GetPath()
        if excel_path != "":
            aaa = excel_path.split("\\")
            a = len(aaa)
            delete_str = aaa[a - 1]
            txt_folder = excel_path.rstrip(delete_str)
            txt_path = txt_folder + txt_path
        else:
            pass

        # print('\n\n\n\n------------------------------------ BIST FAIL Check ------------------------------------')
        cnt_list = []
        slices_n = "-"
        die_n = "-"
        cntv = []
        die_list = []
        with open(txt_path, "r") as f:
            Chip_Info_list = f.readlines()
            index = 0
            for log in Chip_Info_list:
                buffer1 = log.find("unter results")
                get_die = log.find("Mode :")
                if get_die != -1:
                    chk = log
                    tmp = chk.find("Die")
                    if tmp >= 0:
                        die_1 = chk[tmp + 4 : tmp + 5]
                        if die_1 not in die_list:
                            die_list.append(die_1)
                        chk = log[tmp + 5 :]
                        tmp = chk.find("Die")
                        if tmp >= 0:
                            die_2 = chk[tmp + 4 : tmp + 5]
                            if die_2 not in die_list:
                                die_list.append(die_2)
                # info = (index, buffer1, log)  # index : txt row number
                if buffer1 == -1:
                    pass
                else:
                    tmp = log.find("S#")
                    if tmp >= 0:
                        slices_n = log[tmp + 2 : tmp + 3]
                    else:
                        slices_n = "-"
                    tmp = log.find("Die")
                    if tmp >= 0:
                        die_n = log[tmp + 3 : tmp + 4]
                    else:
                        die_n = "-"

                    counter_info = [
                        Chip_Info_list[index + 1],
                        Chip_Info_list[index + 2],
                    ]
                    counter_value_ck = counter_info[0].split(",") + counter_info[
                        1
                    ].split(",")
                    cntv = []
                    for cnt_ck in counter_value_ck:
                        cntv_tmp = cnt_ck.rstrip("\n")
                        cntv += [cntv_tmp.split(" = ")[1]]
                    # print(cntv)
                    cnt_list += [[die_n] + [slices_n] + cntv]
                    if die_n not in die_list:
                        die_list.append(die_n)
                index += 1
        sort_cnt = sorted(cnt_list)
        full_cnt = []
        chk_idx = 0
        if len(sort_cnt) > 0:
            for k in range(len(die_list)):
                for i in range(8):
                    if chk_idx > len(sort_cnt):
                        print(
                            f"check index {chk_idx} > {len(sort_cnt)}, not expected value"
                        )
                    elif chk_idx < len(sort_cnt):
                        chk_item = sort_cnt[chk_idx]
                    if chk_item[0] == die_list[k] and int(chk_item[1]) == i:
                        full_cnt.append(chk_item)
                        chk_idx = chk_idx + 1
                    else:
                        full_cnt.append([die_list[k], f"{i}", "0", "0", "0", "0"])
        out_cnt = [int(x[-i]) for i in range(4, 0, -1) for x in full_cnt]
        # out_cnt = [int(x[-i]) for i in range(4, 0, -1) for x in sort_cnt]
        # cnt_list = numpy.unique(cnt_list).tolist()
        # print(cnt_list)
        return out_cnt

    def txt_log_hwt_check(self, **kargs):
        txt_path = kargs.get("txt_path", "NA")
        mode = kargs.get("mode", "")
        excel_path = self.gui.BIST_Excel_path.GetPath()

        def hwt_1t_check(table, init=0):
            if init == 1:
                left1_list = ["-1"] * 8
                right1_list = ["-1"] * 8
                center1_list = ["-1"] * 8
                win1_list = ["-1"] * 8
                low1_list = ["-1"] * 8
                v_cen1_list = ["-1"] * 8
                v_height1_list = ["-1"] * 8
                high1_list = ["-1"] * 8
            else:
                left1_list = []
                right1_list = []
                center1_list = []
                win1_list = []
                low1_list = []
                v_cen1_list = []
                v_height1_list = []
                high1_list = []
                for i in range(8):
                    buffer_1 = table.split("\n")[i + 4]
                    buffer_1 = buffer_1.replace(" ", "")
                    left1 = buffer_1.split("|")[3]
                    left1_list += [left1]
                    right1 = buffer_1.split("|")[4]
                    right1_list += [right1]
                    win1 = buffer_1.split("|")[6]
                    win1_list += [win1]
                    center1 = buffer_1.split("|")[5]
                    low1 = buffer_1.split("|")[7]
                    low1_list += [low1]
                    high1 = buffer_1.split("|")[8]
                    high1_list += [high1]
                    v_height1 = buffer_1.split("|")[10]
                    v_height1_list += [v_height1]
                    v_cen1 = buffer_1.split("|")[9]
                    if "rpt" in mode:
                        center1 = f"{float(center1) * 100 / 64:2.1f}"
                        if int(float(win1)) == 0:
                            center1 = "N/A"
                        v_cen1 = f"{float(v_cen1) * 100 / 64:2.1f}"
                        if int(float(v_height1)) == 0:
                            v_cen1 = "N/A"
                    center1_list += [center1]
                    v_cen1_list += [v_cen1]
            HW_1T_list = [
                left1_list,
                right1_list,
                center1_list,
                win1_list,
                low1_list,
                high1_list,
                v_cen1_list,
                v_height1_list,
            ]

            return HW_1T_list

        def hwt_check(table):
            check_table = table.split("Read-Training")
            if len(check_table) > 2:
                table_1 = check_table[2]
                table_2 = check_table[1]
                [
                    left1_list,
                    right1_list,
                    center1_list,
                    win1_list,
                    low1_list,
                    high1_list,
                    v_cen1_list,
                    v_height1_list,
                ] = hwt_1t_check(table_1)
                [
                    left2_list,
                    right2_list,
                    center2_list,
                    win2_list,
                    low2_list,
                    high2_list,
                    v_cen2_list,
                    v_height2_list,
                ] = hwt_1t_check(table_2)
            elif len(check_table) > 1:
                if check_table[0].find("train results: RX"):
                    [
                        left1_list,
                        right1_list,
                        center1_list,
                        win1_list,
                        low1_list,
                        high1_list,
                        v_cen1_list,
                        v_height1_list,
                    ] = hwt_1t_check(check_table[1], 1)
                    [
                        left2_list,
                        right2_list,
                        center2_list,
                        win2_list,
                        low2_list,
                        high2_list,
                        v_cen2_list,
                        v_height2_list,
                    ] = hwt_1t_check(check_table[1])
                elif check_table[0].find("train results: TX"):
                    [
                        left1_list,
                        right1_list,
                        center1_list,
                        win1_list,
                        low1_list,
                        high1_list,
                        v_cen1_list,
                        v_height1_list,
                    ] = hwt_1t_check(check_table[1])
                    [
                        left2_list,
                        right2_list,
                        center2_list,
                        win2_list,
                        low2_list,
                        high2_list,
                        v_cen2_list,
                        v_height2_list,
                    ] = hwt_1t_check(check_table[1], 1)
                else:
                    print(f"\033 no find expected training results set tab to -1")
                    [
                        left1_list,
                        right1_list,
                        center1_list,
                        win1_list,
                        low1_list,
                        high1_list,
                        v_cen1_list,
                        v_height1_list,
                    ] = hwt_1t_check(check_table[0], 1)
                    [
                        left2_list,
                        right2_list,
                        center2_list,
                        win2_list,
                        low2_list,
                        high2_list,
                        v_cen2_list,
                        v_height2_list,
                    ] = hwt_1t_check(check_table[0], 1)
            else:
                print(f"\033 no find expected training results set tab to -1")
                [
                    left1_list,
                    right1_list,
                    center1_list,
                    win1_list,
                    low1_list,
                    high1_list,
                    v_cen1_list,
                    v_height1_list,
                ] = hwt_1t_check(check_table[0], 1)
                [
                    left2_list,
                    right2_list,
                    center2_list,
                    win2_list,
                    low2_list,
                    high2_list,
                    v_cen2_list,
                    v_height2_list,
                ] = hwt_1t_check(check_table[0], 1)
            HW_Table = (
                left1_list
                + right1_list
                + left2_list
                + right2_list
                + center1_list
                + win1_list
                + center2_list
                + win2_list
                + low1_list
                + high1_list
                + low2_list
                + high2_list
                + v_cen1_list
                + v_height1_list
                + v_cen2_list
                + v_height2_list
            )
            return HW_Table

        if excel_path != "":
            aaa = excel_path.split("\\")
            a = len(aaa)
            delete_str = aaa[a - 1]
            txt_folder = excel_path.rstrip(delete_str)
            txt_path = txt_folder + txt_path
        else:
            pass
        with open(txt_path, "r") as f:
            table = f.read()
        if len(table.split("Current train results")) > 1:
            table = table.split("Current train results")[1]
            HW_Table = hwt_check(table)
        elif len(table.split("Last train results")) > 1:
            print("Save Last train results in the log")
            HW_Table = hwt_check(table)
        else:
            HW_Table = hwt_check(table)
        return HW_Table

    def txt_log_pwr_check(self, **kargs):
        txt_path = kargs.get("txt_path", "NA")
        last = kargs.get("last", 0)
        excel_path = self.gui.BIST_Excel_path.GetPath()
        if excel_path != "":
            aaa = excel_path.split("\\")
            a = len(aaa)
            delete_str = aaa[a - 1]
            txt_folder = excel_path.rstrip(delete_str)
            txt_path = txt_folder + txt_path
        else:
            pass
        with open(txt_path, "r") as f:
            table = f.read()
        table_C = table.split("Supply Current Measure")
        current_list = []
        for i in range(1, len(table_C)):
            item = table_C[i]
            # each_current = []
            # for item in each_current_o[0:-1]:
            # each_current.append(item.split('/')[0])
            avdd_current = vdd12_1current = vdd12_2current = vddc_current = "Not Test"
            if item.find("AVDD_075: ") >= 0:
                avdd_current = item.split("AVDD_075: ")[1].split("\n")[0]
            if item.find("AVDD12_1: ") >= 0:
                vdd12_1current = item.split("AVDD12_1: ")[1].split(",")[0]
            if item.find("AVDD12_1: ") >= 0:
                vdd12_2current = item.split("AVDD12_2: ")[1].split("\n")[0]
            each_current = [avdd_current, vdd12_1current, vdd12_2current]
            # each_current[-1] = each_current[-1].split('\n')[0]
            # if len(each_current) > 1:
            #     for k in range(len(each_current) - 1, 3):
            #         each_current.append('Not Test')
            current_list = each_current if last == 1 else current_list + each_current
        table_VDDC1 = table.split("VDDC_075 Current ")
        vddc_c_list = []
        for i in range(1, len(table_VDDC1)):
            item = table_VDDC1[i]
            vddc_current = "Not Test"
            if item.find("Value = ") >= 0:
                vddc_current_v = float(item.split("Value = ")[1].split("A")[0])
                vddc_current = f"{vddc_current_v:3.6f}"
        vddc_c_list = [vddc_current] if last == 1 else vddc_c_list + [vddc_current]
        current_list = current_list + vddc_c_list
        return current_list

    def txt_log_full_eye_check(self, **kargs):
        txt_path = kargs.get("txt_path", "NA")
        simple = kargs.get("simple", 0)
        mode = kargs.get("mode", 0)
        offset_apply = kargs.get("off_aply", 1)

        def eye_check(table, offset, eye_info, mode):
            Eye_Table = []
            eye_map = [[0 for _ in range(64)] for _ in range(32)]
            Signal_Eye_Table = [[], [-1] * 32, [-1] * 32, [-1, -1, -1, -1], eye_map]
            for each_eye_t, each_off, each_eye_info in zip(table, offset, eye_info):
                Signal_Eye_Table[3] = each_eye_info.copy()
                Signal_Eye_Table[0] = each_eye_t.split("eye results:\n")[0]
                each_off_text = f"Op{each_off}" if each_off >= 0 else f"On{-each_off}"
                Signal_Eye_Table[0] += each_off_text
                each_eye = each_eye_t.split("eye results:\n")[1]
                each_line = each_eye.split("\n")
                # each_line_o = each_eye.split('\n')
                # each_line = ['0' for _ in range(32)]
                for k in range(32):
                    try:
                        # each_line[k] = each_line_o[k].replace('2', '1')
                        if simple == 1:
                            Signal_Eye_Table[2][k] = (
                                -1
                                if each_line[k].find("1") < 0
                                else 63 - each_line[k].find("1")
                            )
                            Signal_Eye_Table[1][k] = each_line[k][::-1].find("1")
                        else:
                            eye_w = -1
                            e_l = -100
                            e_r = -100
                            for j in range(64):
                                eye_p = (j - each_off) % 64
                                # Signal_Eye_Table[4][k][eye_p] = 1 if each_line[k][63 - j] == '1' else 0
                                Signal_Eye_Table[4][k][eye_p] = int(
                                    each_line[k][63 - j]
                                )
                                if j == 0:
                                    elc = j if each_line[k][63 - j] == "1" else e_l
                                    erc = j if each_line[k][63 - j] == "1" else e_r
                                else:
                                    elc = (
                                        j
                                        if each_line[k][63 - j] == "1"
                                        and each_line[k][64 - j] == "0"
                                        else elc
                                    )
                                    erc = j if each_line[k][63 - j] == "1" else erc
                                ewc = erc - elc + 1
                                if ewc > eye_w:
                                    e_l = elc
                                    e_r = erc
                                    eye_w = ewc
                            if e_r > -100:
                                Signal_Eye_Table[2][k] = e_r - each_off
                                Signal_Eye_Table[1][k] = e_l - each_off
                            else:
                                Signal_Eye_Table[2][k] = e_r
                                Signal_Eye_Table[1][k] = e_l
                    except Exception as e:
                        print(f"From {k} to 31 set to -1 for no information get")
                        Signal_Eye_Table[1][k:32] = [-1] * (32 - k)
                        Signal_Eye_Table[2][k:32] = [-1] * (32 - k)
                        k = k - 1
                        break
                ck_idx = int((k + 1) / 2)
                if mode == 1:
                    if Signal_Eye_Table[2][ck_idx] > -100:
                        Signal_Eye_Table[3][0] = (
                            Signal_Eye_Table[2][ck_idx]
                            - Signal_Eye_Table[1][ck_idx]
                            + 1
                        )
                        Signal_Eye_Table[3][2] = int(
                            (Signal_Eye_Table[2][ck_idx] + Signal_Eye_Table[1][ck_idx])
                            / 2
                        )
                    else:
                        Signal_Eye_Table[3][0] = 0
                    eye_h = -1
                    e_b = -100
                    e_t = -100
                    for j in range(32):
                        if j == 0:
                            ebc = j if Signal_Eye_Table[4][j][32] >= 1 else e_b
                            etc = j if Signal_Eye_Table[4][j][32] >= 1 else e_t
                        else:
                            ebc = (
                                j
                                if Signal_Eye_Table[4][j][32] >= 1
                                and Signal_Eye_Table[4][j - 1][32] == 0
                                else ebc
                            )
                            etc = j if Signal_Eye_Table[4][j][32] >= 1 else etc
                        ehc = etc - ebc + 1
                        if ehc > eye_h:
                            e_b = ebc
                            e_t = etc
                            eye_h = ehc
                    Signal_Eye_Table[3][1] = eye_h if eye_h > -100 else 0
                    Signal_Eye_Table[3][3] = int((e_t + e_b) / 2) if eye_h > -100 else 0
                Eye_Table.append(
                    [
                        Signal_Eye_Table[0],
                        Signal_Eye_Table[1].copy(),
                        Signal_Eye_Table[2].copy(),
                        Signal_Eye_Table[3].copy(),
                        copy.deepcopy(Signal_Eye_Table[4]),
                    ]
                )
            return Eye_Table

        excel_path = self.gui.shmoo_load_path.GetPath()
        if excel_path != "":
            aaa = excel_path.split("\\")
            a = len(aaa)
            delete_str = aaa[a - 1]
            txt_folder = excel_path.rstrip(delete_str)
            txt_path = txt_folder + txt_path
        else:
            pass
        with open(txt_path, "r") as f:
            table = f.read()
        if len(table.split("eye results:")) > 1:
            offset_list = []
            eye_info_list = []
            table_off_all = table.split("print die")
            for table_off in table_off_all:
                if table_off.find("Read-Training") >= 0:
                    table_off = table_off.split("Read-Training")[1]
                    table_off = table_off.split("\n")[4:12]
                    for t in range(8):
                        w_c = int(table_off[t].split("|")[5])
                        if w_c >= 0:
                            offset_list.append(w_c - 31)
                        e_w = float(table_off[t].split("|")[6])
                        if len(table_off[t].split("|")) >= 10:
                            h_c = int(table_off[t].split("|")[9])
                            e_h = float(table_off[t].split("|")[10])
                        if e_w >= 0:
                            eye_info_list.append([e_w, e_h, w_c, h_c])
            table_eye = table.split("print die")[1:]
            if offset_apply == 1:
                Eye_Table = eye_check(table_eye, offset_list, eye_info_list, mode)
            else:
                Eye_Table = eye_check(
                    table_eye, [0] * len(offset_list), eye_info_list, mode
                )
        else:
            print("No Full Eye results find")
            Eye_Table = eye_check([], [], [], [])
        return Eye_Table

    def Graph_TX_Eye_and_Jitter(self):
        font = Font(
            "Arial", size=10, bold=False, italic=False, strike=False, color="000000"
        )
        fig, ax = plt.subplots()
        fig.set_figheight(5)
        fig.set_figwidth(10)
        ax.spines["top"].set_visible(True)
        ax.spines["right"].set_visible(False)
        ax.spines["bottom"].set_visible(True)
        ax.spines["left"].set_visible(True)
        Graph_ScaleSize = 15

        time.sleep(0.2)
        doc = Document(r"Test Report.docx")
        # time.sleep(0.1)
        tables = doc.tables
        # print(table_num)
        table = tables[table_num]
        # print(table)

        # Edit test value graph
        for i in range(TestItem_number):
            List = []
            for row in range(5):
                # print(text)
                my_list = table.cell(i + 1, row + 1).text  # (row, cols)
                # print(my_list)
                if my_list == "NA":
                    my_list = -1
                my_list = [float(my_list)]
                List += my_list
                # print(List)
                string = table.cell(i + 1, 0).text  # (row, cols)
            x = ["LV - LT", "HV - LT", "NV - NT", "LV - HT", "HV - HT"]
            plt.plot(x, List, "o-", label=string, linewidth=3, markersize=8)

        # specification Graph
        for i in range(1):
            try:
                my_list = table.cell(TestItem_number + 1, 1).text
                # print(my_list)
                # Specification define range
                # print(Graph_title)
                if "to" in TestSpec_val:
                    List = []
                    for i in range(5):
                        my_list = table.cell(TestItem_number + 1, 1).text  # (row, cols)
                        # print(my_list)
                        List += [float(my_list)]
                    # print(List)
                    Specification_Min = List
                    List = []
                    for i in range(5):
                        my_list = table.cell(TestItem_number + 2, 1).text  # (row, cols)
                        # print(my_list)
                        List += [float(my_list)]
                    # print(List)
                    Specification_Max = List
                    x = ["LV - LT", "HV - LT", "NV - NT", "LV - HT", "HV - HT"]
                    plt.plot(
                        x,
                        Specification_Max,
                        "D--",
                        label="Spec. Min",
                        linewidth=3,
                        markersize=8,
                        c="r",
                    )
                    plt.plot(
                        x,
                        Specification_Min,
                        "D--",
                        label="Spec. Max",
                        linewidth=3,
                        markersize=8,
                        c="r",
                    )
                elif "Jitter" in Graph_title:
                    # Specification define max
                    List = []
                    for i in range(5):
                        my_list = table.cell(TestItem_number + 1, 1).text  # (row, cols)
                        if my_list == "NA":
                            my_list = -1
                        elif my_list == "NA\n":
                            my_list = -1
                        # print(my_list)
                        List += [float(my_list)]
                    # print(List)
                    Specification = List
                    x = ["LV - LT", "HV - LT", "NV - NT", "LV - HT", "HV - HT"]
                    plt.plot(
                        x,
                        Specification,
                        "D--",
                        label="Spec. Max",
                        linewidth=3,
                        markersize=8,
                        c="r",
                    )
                else:
                    # Specification define min/max
                    List = []
                    for i in range(5):
                        my_list = table.cell(TestItem_number + 1, 1).text  # (row, cols)
                        if my_list == "NA":
                            my_list = -1
                        elif my_list == "NA\n":
                            my_list = -1
                        # print(my_list)
                        List += [float(my_list)]
                    # print(List)
                    Specification = List
                    x = ["LV - LT", "HV - LT", "NV - NT", "LV - HT", "HV - HT"]
                    plt.plot(
                        x,
                        Specification,
                        "D--",
                        label="Spec. Min",
                        linewidth=3,
                        markersize=8,
                        c="r",
                    )
            except Exception:
                break

        font = {"family": "Arial", "weight": "normal", "size": 10}
        plt.title(
            Intreface_name + TestFunction_val + "\n", font, fontsize=Graph_ScaleSize
        )
        # plt.xlabel("XX")
        plt.xticks(fontsize=Graph_ScaleSize)
        plt.yticks(fontsize=Graph_ScaleSize)
        plt.ylabel(Graph_title, font, fontsize=Graph_ScaleSize)
        plt.legend(
            loc="upper left",
            bbox_to_anchor=(1, 1),
            frameon=False,
            borderaxespad=0,
            fontsize=Graph_ScaleSize,
        )
        plt.grid(axis="y", c="black")
        fig.tight_layout()  # Auto Set graph and note range

        plt.savefig("Graph.png")

        img = cv.imread("Graph.png")
        # cv.namedWindow('Graph.png', cv.WINDOW_NORMAL)
        # cv.resizeWindow('Graph.png', 500, 500)
        img = cv.copyMakeBorder(
            img, 1, 1, 1, 1, cv.BORDER_CONSTANT, value=[1, 1, 1]
        )  # 添加边框
        cv.imwrite("Graph.png", img)
        # cv.imshow('Graph.png', img)
        # cv.waitKey(0)
        # cv.destroyAllWindows()

        plt.cla()
        plt.close("all")
        # plt.show()

    def Graph_kvco(self, **kargs):
        # Create Grap Flow :
        # 1) Test result will save to "Test_Graph_Log.txt"
        # 2) python tools will read "Test_Graph_Log.txt" data to create test graph

        # Test_Graph_Log.txt (Sample)
        # File Name, Chip Number, Chip Version, Voltage(V), Temterture(Degree), Band, Reslut, Reslut, Reslut, Reslut, Reslut, Reslut, Reslut, Reslut, Reslut, Reslut, Reslut, Reslut, Reslut, Reslut, Reslut, Reslut, Reslut, Reslut, Reslut, Reslut, Reslut, Reslut, Reslut, Reslut, Reslut, Reslut, Reslut, Reslut, Reslut, Reslut, Reslut, Reslut, Reslut, Reslut, Reslut, Reslut, Reslut, Reslut, Reslut, Reslut, Reslut, Reslut, Reslut, Reslut, Reslut, Reslut, Reslut, Reslut, Reslut, Reslut, Reslut, Reslut, Reslut, Reslut, Reslut, Reslut, Reslut, Reslut, Reslut, Reslut, Reslut, Reslut, Reslut, Reslut, Reslut, Reslut, Reslut, Reslut, Reslut, Reslut, Reslut, Reslut, Reslut, Reslut, Reslut, Reslut, Reslut, Reslut
        # NV - NT, TT01, EVE008A, 0.9, 25, 111, 1.78 / 30400.00, 1.78 / 30464.00, 1.78 / 30528.00, 1.78 / 30592.00, 1.78 / 30656.00, 1.78 / 30720.00, 1.78 / 30784.00, 1.78 / 30848.00, 1.78 / 30912.00, 1.78 / 30976.00, 1.78 / 31040.00, 1.78 / 31104.00, 1.78 / 31168.00, 1.78 / 31232.00, 1.78 / 31296.00, 1.78 / 31360.00, 1.78 / 31424.00, 1.78 / 31488.00, 1.78 / 31552.00, 1.78 / 31616.00, 1.78 / 31680.00, 1.78 / 31744.00, 1.78 / 31808.00, 1.78 / 31872.00, 1.78 / 31936.00, 1.72 / 32000.00, 1.59 / 32064.00, 1.50 / 32128.00, 1.45 / 32192.00, 1.41 / 32256.00, 1.38 / 32320.00, 1.36 / 32384.00, 1.34 / 32448.00, 1.32 / 32512.00, 1.30 / 32576.00, 1.29 / 32640.00, 1.27 / 32704.00, 1.25 / 32768.00, 1.24 / 32832.00, 1.23 / 32896.00, 1.21 / 32960.00, 1.20 / 33024.00, 1.19 / 33088.00, 1.17 / 33152.00, 1.16 / 33216.00, 1.15 / 33280.00, 1.13 / 33344.00, 1.12 / 33408.00, 1.10 / 33472.00, 1.09 / 33536.00, 1.07 / 33600.00, 1.06 / 33664.00, 1.04 / 33728.00, 1.03 / 33792.00, 1.01 / 33856.00, 0.99 / 33920.00, 0.98 / 33984.00, 0.95 / 34048.00, 0.93 / 34112.00, 0.89 / 34176.00, 0.85 / 34240.00, 0.79 / 34304.00, 0.72 / 34368.00, 0.63 / 34432.00, 0.52 / 34496.00, 0.38 / 34560.00, 0.22 / 34624.00, 0.02 / 34688.00, 0.01 / 34752.00, 0.01 / 34816.00, 0.01 / 34880.00, 0.01 / 34944.00, 0.01 / 35008.00, 0.01 / 35072.00, 0.01 / 35136.00, 0.01 / 35200.00, x, Test1
        # NV - NT, TT01, EVE008A, 0.9, 125, 111, 1.78 / 30400.00, 1.78 / 30464.00, 1.78 / 30528.00, 1.78 / 30592.00, 1.78 / 30656.00, 1.78 / 30720.00, 1.78 / 30784.00, 1.78 / 30848.00, 1.78 / 30912.00, 1.78 / 30976.00, 1.78 / 31040.00, 1.78 / 31104.00, 1.78 / 31168.00, 1.78 / 31232.00, 1.78 / 31296.00, 1.78 / 31360.00, 1.78 / 31424.00, 1.78 / 31488.00, 1.78 / 31552.00, 1.78 / 31616.00, 1.78 / 31680.00, 1.78 / 31744.00, 1.78 / 31808.00, 1.78 / 31872.00, 1.78 / 31936.00, 1.72 / 32000.00, 1.59 / 32064.00, 1.50 / 32128.00, 1.45 / 32192.00, 1.41 / 32256.00, 1.38 / 32320.00, 1.36 / 32384.00, 1.34 / 32448.00, 1.32 / 32512.00, 1.30 / 32576.00, 1.29 / 32640.00, 1.27 / 32704.00, 1.25 / 32768.00, 1.24 / 32832.00, 1.23 / 32896.00, 1.21 / 32960.00, 1.20 / 33024.00, 1.19 / 33088.00, 1.17 / 33152.00, 1.16 / 33216.00, 1.15 / 33280.00, 1.13 / 33344.00, 1.12 / 33408.00, 1.10 / 33472.00, 1.09 / 33536.00, 1.07 / 33600.00, 1.06 / 33664.00, 1.04 / 33728.00, 1.03 / 33792.00, 1.01 / 33856.00, 0.99 / 33920.00, 0.98 / 33984.00, 0.95 / 34048.00, 0.93 / 34112.00, 0.89 / 34176.00, 0.85 / 34240.00, 0.79 / 34304.00, 0.72 / 34368.00, 0.63 / 34432.00, 0.52 / 34496.00, 0.38 / 34560.00, 0.22 / 34624.00, 0.02 / 34688.00, 0.01 / 34752.00, 0.01 / 34816.00, 0.01 / 34880.00, 0.01 / 34944.00, 0.01 / 35008.00, 0.01 / 35072.00, 0.01 / 35136.00, 0.05 / 35200.00, x, Test2

        png_name = self.gui.vco_graph_name.Value

        Txt_Data = open(r"TestTools(EY0012A)\Test_Graph_Log.txt")
        txt_lines_num = len(Txt_Data.readlines()) - 1
        # print(txt_lines_num)
        Txt_Data = open(r"TestTools(EY0012A)\Test_Graph_Log.txt")
        text = []
        for line in Txt_Data:
            text.append(line)
        del text[0]

        fig = plt.figure(figsize=(25, 20))
        ax = fig.add_axes([0.1, 0.1, 0.6, 0.8])  # Graph picture position and size
        ax.spines["right"].set_visible(True)
        ax.spines["top"].set_visible(True)
        # plt.tight_layout()

        Color_array = []
        List0 = []

        for i in range(txt_lines_num):  # for i in range(txt_lines_num):
            list0 = text[i].split(",")
            # print(list0)
            Color_array = [
                "b",
                "c",
                "k",
                "g",
                "m",
                "r",
                "y",
                "tab:blue",
                "tab:orange",
                "tab:green",
                "tab:red",
                "tab:purple",
                "tab:brown",
                "tab:pink",
                "tab:gray",
                "tab:olive",
                "tab:cyan",
                "b",
                "c",
                "k",
                "g",
                "m",
                "r",
                "y",
                "tab:blue",
                "tab:orange",
                "tab:green",
                "tab:red",
                "tab:purple",
                "tab:brown",
                "tab:pink",
                "tab:gray",
                "tab:olive",
                "tab:cyan",
                "b",
                "c",
                "k",
                "g",
                "m",
                "r",
                "y",
                "tab:blue",
                "tab:orange",
                "tab:green",
                "tab:red",
                "tab:purple",
                "tab:brown",
                "tab:pink",
                "tab:gray",
                "tab:olive",
                "tab:cyan",
                "b",
                "c",
                "k",
                "g",
                "m",
                "r",
                "y",
                "tab:blue",
                "tab:orange",
                "tab:green",
                "tab:red",
                "tab:purple",
                "tab:brown",
                "tab:pink",
                "tab:gray",
                "tab:olive",
                "tab:cyan",
                "b",
                "c",
                "k",
                "g",
                "m",
                "r",
                "y",
                "tab:blue",
                "tab:orange",
                "tab:green",
                "tab:red",
                "tab:purple",
                "tab:brown",
                "tab:pink",
                "tab:gray",
                "tab:olive",
                "tab:cyan",
                "b",
                "c",
                "k",
                "g",
                "m",
                "r",
                "y",
                "tab:blue",
                "tab:orange",
                "tab:green",
                "tab:red",
                "tab:purple",
                "tab:brown",
                "tab:pink",
                "tab:gray",
                "tab:olive",
                "tab:cyan",
                "b",
                "c",
                "k",
                "g",
                "m",
                "r",
                "y",
                "tab:blue",
                "tab:orange",
                "tab:green",
                "tab:red",
                "tab:purple",
                "tab:brown",
                "tab:pink",
                "tab:gray",
                "tab:olive",
                "tab:cyan",
                "b",
                "c",
                "k",
                "g",
                "m",
                "r",
                "y",
                "tab:blue",
                "tab:orange",
                "tab:green",
                "tab:red",
                "tab:purple",
                "tab:brown",
                "tab:pink",
                "tab:gray",
                "tab:olive",
                "tab:cyan",
                "b",
                "c",
                "k",
                "g",
                "m",
                "r",
                "y",
                "tab:blue",
                "tab:orange",
                "tab:green",
                "tab:red",
                "tab:purple",
                "tab:brown",
                "tab:pink",
                "tab:gray",
                "tab:olive",
                "tab:cyan",
                "b",
                "c",
                "k",
                "g",
                "m",
                "r",
                "y",
                "tab:blue",
                "tab:orange",
                "tab:green",
                "tab:red",
                "tab:purple",
                "tab:brown",
                "tab:pink",
                "tab:gray",
                "tab:olive",
                "tab:cyan",
                "b",
                "c",
                "k",
                "g",
                "m",
                "r",
                "y",
                "tab:blue",
                "tab:orange",
                "tab:green",
                "tab:red",
                "tab:purple",
                "tab:brown",
                "tab:pink",
                "tab:gray",
                "tab:olive",
                "tab:cyan",
                "b",
                "c",
                "k",
                "g",
                "m",
                "r",
                "y",
                "tab:blue",
                "tab:orange",
                "tab:green",
                "tab:red",
                "tab:purple",
                "tab:brown",
                "tab:pink",
                "tab:gray",
                "tab:olive",
                "tab:cyan",
                "b",
                "c",
                "k",
                "g",
                "m",
                "r",
                "y",
                "tab:blue",
                "tab:orange",
                "tab:green",
                "tab:red",
                "tab:purple",
                "tab:brown",
                "tab:pink",
                "tab:gray",
                "tab:olive",
                "tab:cyan",
                "b",
                "c",
                "k",
                "g",
                "m",
                "r",
                "y",
                "tab:blue",
                "tab:orange",
                "tab:green",
                "tab:red",
                "tab:purple",
                "tab:brown",
                "tab:pink",
                "tab:gray",
                "tab:olive",
                "tab:cyan",
                "b",
                "c",
                "k",
                "g",
                "m",
                "r",
                "y",
                "tab:blue",
                "tab:orange",
                "tab:green",
                "tab:red",
                "tab:purple",
                "tab:brown",
                "tab:pink",
                "tab:gray",
                "tab:olive",
                "tab:cyan",
                "b",
                "c",
                "k",
                "g",
                "m",
                "r",
                "y",
                "tab:blue",
                "tab:orange",
                "tab:green",
                "tab:red",
                "tab:purple",
                "tab:brown",
                "tab:pink",
                "tab:gray",
                "tab:olive",
                "tab:cyan",
                "b",
                "c",
                "k",
                "g",
                "m",
                "r",
                "y",
                "tab:blue",
                "tab:orange",
                "tab:green",
                "tab:red",
                "tab:purple",
                "tab:brown",
                "tab:pink",
                "tab:gray",
                "tab:olive",
                "tab:cyan",
                "b",
                "c",
                "k",
                "g",
                "m",
                "r",
                "y",
                "tab:blue",
                "tab:orange",
                "tab:green",
                "tab:red",
                "tab:purple",
                "tab:brown",
                "tab:pink",
                "tab:gray",
                "tab:olive",
                "tab:cyan",
                "b",
                "c",
                "k",
                "g",
                "m",
                "r",
                "y",
                "tab:blue",
                "tab:orange",
                "tab:green",
                "tab:red",
                "tab:purple",
                "tab:brown",
                "tab:pink",
                "tab:gray",
                "tab:olive",
                "tab:cyan",
                "b",
                "c",
                "k",
                "g",
                "m",
                "r",
                "y",
                "tab:blue",
                "tab:orange",
                "tab:green",
                "tab:red",
                "tab:purple",
                "tab:brown",
                "tab:pink",
                "tab:gray",
                "tab:olive",
                "tab:cyan",
                "b",
                "c",
                "k",
                "g",
                "m",
                "r",
                "y",
                "tab:blue",
                "tab:orange",
                "tab:green",
                "tab:red",
                "tab:purple",
                "tab:brown",
                "tab:pink",
                "tab:gray",
                "tab:olive",
                "tab:cyan",
            ]
            Line_color = Color_array[i]
            # Read kvco Line_Name for graph lable use
            kvco_line_name = (list0[-1]).split("\n")
            kvco_line_name = kvco_line_name[0]
            # print(kvco_line_name)

            # Read kvco color for graph use
            kvco_line_color = (list0[-2]).split("\n")
            kvco_line_color = str(kvco_line_color[0])
            # print(kvco_line_color)
            if kvco_line_color == "Auto":
                kvco_line_color = "black"

            else:
                kvco_line_color = "blue"
            # print(kvco_line_color)

            # Read kvco test result
            del list0[0:1]
            # del list0[0:6]
            # print(list0)
            del list0[-1]
            # print(list0)
            del list0[-1]
            # print(list0)

            list_len = len(list0)
            # print(list_len)

            # create graph
            List1 = []
            List2 = []
            List3 = []
            for i in range(list_len):
                Freq_Amp = list0[i].split("/")
                my_list = Freq_Amp[0]
                List1 += [my_list]
            List1 = [x.replace("None", "") for x in List1]
            List1 = [i for i in List1 if i]
            List1 = [float(x) for x in List1]
            List1 = List1[::-1]

            for i in range(list_len):
                try:
                    Freq_Amp = list0[i].split("/")
                    my_list = Freq_Amp[1]
                    List2 += [my_list]
                    List2 = [x.replace("None", "") for x in List2]
                    List2 = [i for i in List2 if i]
                except:
                    break
            List2 = [x for x in List2 if x is not None]
            List2 = [float(x) for x in List2]
            List2 = List2[::-1]

            plt.plot(
                List1,
                List2,
                "o-",
                label=kvco_line_name,
                linewidth=1,
                markersize=3,
                c=Line_color,
            )
            plt.rcParams["font.family"] = "Arial"
            plt.rcParams["font.size"] = "7"  # lable size

        plt.tick_params(labelsize=6)

        # ax.set_yticks(List2)
        # ax.set_yticklabels(List2, rotation=0, ha='right')
        # plt.ylim(Freq_Min, Freq_Max)
        # plt.xlim(20, 100)

        plt.minorticks_on()
        plt.tick_params(which="both", direction="in")

        # mylist = []
        # for i in list0:
        #     mylist += i
        # print(mylist)

        font = {"family": "Arial", "weight": "normal", "size": 8}
        plt.title("KVCO Graph", font)
        plt.ylabel("Frequency * 4 * sel_div (Unit : MHz)\n", font)
        plt.xlabel("\nVtune (Unit : V)\n", font)

        plt.tick_params(labelsize=8)
        labels = ax.get_xticklabels() + ax.get_yticklabels()
        [label.set_fontname("Arial") for label in labels]

        plt.legend(
            bbox_to_anchor=(1.01, 1), loc="upper left", borderaxespad=0.0
        )  # Show note lable

        file_name = "Test Graph(EY0012A)/" + png_name + ".png"
        plt.savefig(file_name)

        # plt.show()
        print("Graph_kvco is done")

    def Graph_slewRate(self):
        # Create Grap Flow :
        # 1) Test result will save to "Test_Report_Log.txt",
        # 2) The Test_Report_Log.txt file have scope waveform .csv file path address
        # 3) Python tools will call cav path address and create slewrate graph picture
        # 4) Read Test_Report_Log.txt--> Graph_File_Name , save graph picture

        # Test_Report_Log.txt
        # Graph_File_Name: Test
        # C:\Users\Vince\Desktop\Python\Test Report\Test Waveform\Slew_Reg = 00.txt
        # C:\Users\Vince\Desktop\Python\Test Report\Test Waveform\Slew_Reg = 01.txt
        # C:\Users\Vince\Desktop\Python\Test Report\Test Waveform\Slew_Reg = 02.txt
        # C:\Users\Vince\Desktop\Python\Test Report\Test Waveform\Slew_Reg = 03.txt

        fig = plt.figure()
        ax = fig.add_axes([0.25, 0.25, 0.5, 0.5])  # Graph picture position and size
        ax.spines["right"].set_visible(True)
        ax.spines["top"].set_visible(True)
        plt.tight_layout()

        Txt_Data = open(r"Test_Report_Log.txt")
        csv_list = Txt_Data.readlines()  # read Test_Report_Log.txt data

        # print(len(csv_list))
        for csv in range(len(csv_list) - 1):
            try:
                # print(csv)
                graph_color_array = [
                    "k",
                    "r",
                    "c",
                    "b",
                    "g",
                    "c",
                    "m",
                    "y",
                    "tan",
                ]  # define graph line color
                Txt_Data = open(r"Test_Report_Log.txt")  # read Test_Report_Log.txt data
                csv_list = Txt_Data.read().splitlines()  # delete list "\n" iten
                print(csv_list)

                List1 = []
                List2 = []
                csv_list = (csv_list[csv + 1]).split("\n")  # delete string "\n"
                # print(csv_list)
                csv_list = csv_list[0]

                csv_list = csv_list.split(";")
                # print(csv_list)
                Graph_lable_name = csv_list[1]  # graph lable rigistor name
                # print(Graph_lable_name)
                csv_list = csv_list[0]  # csv path
                # print(csv_list)

                Txt_Data = open(csv_list, "r")  # read time/amp list
                csv_list = Txt_Data.readlines()
                for i in range(9999):
                    try:
                        # print(i)
                        csv_val = csv_list[i + 20]
                        csv_val = csv_val.split(",")
                        # print(csv_val)
                        my_list = csv_val[0]
                        List1 += [float(my_list)]
                        # print(List1)

                        csv_vol_val = csv_val[1].split("\n")
                        csv_vol_val = (csv_vol_val[0]).split(" ")
                        my_list = csv_vol_val[1]
                        # print(my_list)
                        List2 += [float(my_list)]
                        # print(List2)
                        # print('\n')
                    except Exception:
                        break
                # print(i)
                Txt_Data.close()

                # edit line name
                Txt_Data = open(r"Test_Report_Log.txt")
                Line_name = Txt_Data.readlines()
                Line_name = Line_name[csv + 1].split("Waveform\\")
                Line_name = Line_name[1].split("\n")  # split string
                # print(Line_name)
                Line_name = Line_name[0].replace(".txt", "")
                # print(Line_name)
                plt.plot(
                    List1,
                    List2,
                    "-",
                    label=Graph_lable_name,
                    linewidth=1,
                    markersize=3,
                    c=graph_color_array[csv + 1],
                )
            except Exception:
                break

        font = {"family": "Arial", "weight": "normal", "size": 10}
        plt.rcParams["font.family"] = "Arial"
        plt.rcParams["font.size"] = "10"
        plt.tick_params(labelsize=6)
        plt.minorticks_on()
        plt.tick_params(which="both", direction="in")
        plt.ylabel("V peak to peak(Unit : MHz)\n", font)
        plt.xlabel("\n Time(Unit : V)\n", font)
        plt.tick_params(labelsize=8)
        labels = ax.get_xticklabels() + ax.get_yticklabels()
        [label.set_fontname("Arial") for label in labels]
        plt.legend(
            bbox_to_anchor=(1.01, 1), loc="upper left", borderaxespad=0.1
        )  # Show note lable

        Txt_Data = open(r"Test_Report_Log.txt")
        Pic_name = Txt_Data.readlines()
        Pic_name = Pic_name[0].split(":")
        Pic_name = Pic_name[1].split("\n")
        Pic_name = Pic_name[0].split(" ")
        Pic_name = Pic_name[1]
        # print(Pic_name)

        Txt_Data = open(r"Test_Report_Log.txt")
        Pic_path = Txt_Data.readlines()
        Pic_path = Pic_path[1].split("Test Waveform\\")
        Pic_path = Pic_path[0] + "Test Waveform/"
        # print(Pic_path)
        plt.title(Pic_name, font)
        plt.savefig(Pic_path + Pic_name + ".png")

        # print(Pic_name + '.png')
        plt.show()

        # lines = text_file.read().split('Registor')
        # print(len(lines))
        # text_file.close()

    def Graph_shmoo_voltage(self, **kargs):
        title = kargs.get("title", "")
        """
        Shmoo test log sample
        Analog Voltage,Digital Voltage,Pass/Fail
        0.800,0.800,1
        0.800,1.000,1
        1.000,0.800,1
        1.000,1.000,1
        """
        sel_avdd = self.gui.shmoo_VDDC_voltage.Value
        font = {"family": "Arial", "weight": "normal", "size": 10}
        df = pd.read_csv("TestTools(EY0012A)/Buffer.txt")
        df_fail = df.loc[(df["Pass/Fail"] == 0)]
        df_fail.reset_index(drop=True, inplace=True)
        df_pass = df.loc[(df["Pass/Fail"] == 1)]
        df_pass.reset_index(drop=True, inplace=True)
        fail_pts = (df_fail["Analog Voltage"], df_fail["Digital Voltage"])
        pass_pts = (df_pass["Analog Voltage"], df_pass["Digital Voltage"])

        data = (pass_pts, fail_pts)
        colors = ("lime", "red")
        groups = ("Pass", "Fail")

        # Create plot
        for data, color, group in zip(data, colors, groups):
            x, y = data
            plt.scatter(
                x,
                y,
                alpha=1,
                c=color,
                edgecolors="black",
                s=250,
                label=group,
                marker="s",
            )

        # plt.title('Voltage Shmoo', fontdict=font)
        plt.title(title + "_voltage shmoo test", font)
        plt.grid(True, which="both")

        shift = 0.02 * int((float(sel_avdd) - 0.75) / 0.02)  # use 0.75 as default
        # plt.xticks(np.arange(min(x), max(x) + 1, 0.02))
        # plt.yticks(np.arange(min(x), max(x) + 1, 0.02))
        plt.xticks(np.arange(0.65, 0.89, 0.02))
        plt.yticks(np.arange(0.65 + shift, 0.89 + shift, 0.02))
        plt.xlim(0.57, 0.89)
        plt.ylim(0.63 + shift, 0.89 + shift)
        # plt.yticks(np.arange(0.57, 0.83, 0.02))
        # plt.ylim(0.55, 0.83)

        plt.xlabel("Digital Voltage (Unit : V)", font)
        plt.ylabel("Analog Voltage (Unit : V)", font)
        plt.legend(loc=2, frameon=True)
        plt.tight_layout()
        plt.savefig("Test Graph(EY0012A)/Shmoo_Voltage_" + title + ".png")
        plt.show()

    def Graph_shmoo_vref(self, **kargs):
        path = kargs.get("path", self.gui.shmoo_load_path.GetPath())
        # NV_voltage = kargs.get('NV_voltage', 0.75)
        # sheet = kargs.get('sheet_type', '')
        # type = kargs.get('type', 0.75)

        work_book = openpyxl.load_workbook(path)
        fp_name = path.split("\\")
        fname = fp_name[-1]
        path = path.rstrip(fname)
        sheet = work_book[self.gui.shmoo_type.Value]
        sel_chip = self.gui.shmoo_chip_number.Value
        sel_mode = self.gui.shmoo_chip_mode.Value
        sel_temper = self.gui.shmoo_test_temperature.Value
        sel_avdd = self.gui.shmoo_VDDC_voltage.Value

        chip_col = None
        temper_col = None
        avdd_col = None
        mode_col = None
        vref_col = None
        win_left_col = []
        title_list = []
        for col in sheet["4:4"]:
            cell_v = sheet.cell(row=col.row, column=col.column).value
            if cell_v is not None:
                title_list.append(cell_v)
                if "TRAIN" in cell_v.upper() and "WIN_LEFT" in cell_v.upper():
                    train_win_col = col.column
                    win_left_col.append(train_win_col)
        for col in sheet["5:5"]:
            cell_v = sheet.cell(row=col.row, column=col.column).value
            if cell_v is not None and type(cell_v) == str:
                title_list.append(cell_v)
                if "CHIP NUM" in cell_v.upper():
                    chip_col = col.column
                if "TEMPERATURE" in cell_v.upper():
                    temper_col = col.column
                if "CHANNEL1" in cell_v.upper() and "VOLTAGE" in cell_v.upper():
                    avdd_col = col.column
                if "MODE" in cell_v.upper():
                    mode_col = col.column
                if "REGISTER_1" in cell_v.upper():
                    vref_col = col.column

        vref_list = []
        check_row = []
        for z in range(6, sheet.max_row + 1):
            if (
                sel_chip == sheet.cell(row=z, column=chip_col).value
                and sel_mode == sheet.cell(row=z, column=mode_col).value
                and sel_temper == str(sheet.cell(row=z, column=temper_col).value)
            ):
                if float(sel_avdd) == sheet.cell(row=z, column=avdd_col).value:
                    vref_get = sheet.cell(row=z, column=vref_col).value
                    if "0X" in vref_get.upper():
                        base = 16
                    else:
                        base = 10
                    vref_list.append(int(vref_get, base))
                    check_row.append(z)
        # print('check_row=',check_row)
        if len(check_row) == 0:
            print(f"no data for plot")
            return
        # "win_left Excel's col num"_"Die and slice name"
        for slice_sel in [
            "Die1-Slice-0",
            "Die1-Slice-1",
            "Die1-Slice-2",
            "Die1-Slice-3",
            "Die1-Slice-4",
            "Die1-Slice-5",
            "Die1-Slice-6",
            "Die1-Slice-7",
            "Die0-Slice-0",
            "Die0-Slice-1",
            "Die0-Slice-2",
            "Die0-Slice-3",
            "Die0-Slice-4",
            "Die0-Slice-5",
            "Die0-Slice-6",
            "Die0-Slice-7",
        ]:
            slice_n = int((slice_sel.split("-"))[2])
            if "Die1" in slice_sel:
                die_chk = 0
            else:
                die_chk = 1
            f_Title = f"{sel_chip}_{sel_mode}_T{sel_temper}_{int(float(sel_avdd) * 1000)}mV_{slice_sel}"
            Title = f"{slice_sel}"

            slice_num = win_left_col[die_chk] + slice_n
            self.Graph_shmoo_vref_plot(
                check_row=check_row,
                sheet=sheet,
                sel_avdd=sel_avdd,
                vref_list=vref_list,
                slice_num=slice_num,
                path=path,
                Title=Title,
                f_Title=f_Title,
            )

    def Graph_shmoo_vref_train(self, **kargs):
        path = kargs.get("path", self.gui.shmoo_load_path.GetPath())
        xls_path = path
        # NV_voltage = kargs.get('NV_voltage', 0.75)
        # sheet = kargs.get('sheet_type', '')
        # type = kargs.get('type', 0.75)
        demo_mode = 0
        detail_m = 0
        start_row = 6
        eye_sn = [[[[0] * 1 for _ in range(8)] for _ in range(4)] for _ in range(4)]
        work_book = openpyxl.load_workbook(path)
        fp_name = path.split("\\")
        fname = fp_name[-1]
        path = path.rstrip(fname)
        # sheet = work_book[self.gui.shmoo_type.Value]
        font = {"family": "Arial", "weight": "normal", "size": 10}
        off_aply = 1
        if "@" in self.gui.shmoo_type.Value:
            if "DEMO" in self.gui.shmoo_type.Value.split("@")[1].upper():
                sheet = work_book["Manual_Mode"]
                demo_mode = 1
            elif "CEN" in self.gui.shmoo_type.Value.split("@")[1].upper():
                if "CEN0" in self.gui.shmoo_type.Value.split("@")[1].upper():
                    off_aply = 0
                sheet = work_book["Manual_Mode"]
                detail_m = 1
            else:
                sheet = work_book[self.gui.shmoo_type.Value.split("@")[0]]
                if "CEN" in self.gui.shmoo_type.Value.split("@")[1].upper():
                    if "CEN0" in self.gui.shmoo_type.Value.split("@")[1].upper():
                        off_aply = 0
                    detail_m = 1
        else:
            sheet = work_book["Manual_Mode"]
        ws = sheet
        sel_chip = self.gui.shmoo_chip_number.Value
        sel_mode = self.gui.shmoo_chip_mode.Value
        sel_temper = self.gui.shmoo_test_temperature.Value
        sel_avdd = self.gui.shmoo_VDDC_voltage.Value

        chip_col = None
        temper_col = None
        avdd_col = None
        mode_col = None
        vref_col = None
        log_col = None
        win_left_col = []
        title_list = []
        font = Font(
            "Arial", size=10, bold=False, italic=False, strike=False, color="000000"
        )
        border = Border(
            left=Side(border_style="thin", color="000000"),
            right=Side(border_style="thin", color="000000"),
            top=Side(border_style="thin", color="000000"),
            bottom=Side(border_style="thin", color="000000"),
        )

        align = Alignment(horizontal="center", vertical="center", wrap_text=True)
        # for col in sheet['4:4']:
        #     cell_v = sheet.cell(row=col.row, column=col.column).value
        #     if cell_v is not None:
        #         title_list.append(cell_v)
        #         if 'TRAIN' in cell_v.upper() and 'WIN_LEFT' in cell_v.upper():
        #             train_win_col = col.column
        #             win_left_col.append(train_win_col)
        #
        for col in sheet["5:5"]:
            cell_v = sheet.cell(row=col.row, column=col.column).value
            if cell_v is not None and type(cell_v) == str:
                title_list.append(cell_v)
                if "CHIP NUM" in cell_v.upper():
                    chip_col = col.column
                if "TEMPERATURE" in cell_v.upper():
                    temper_col = col.column
                # if 'CHANNEL1' in cell_v.upper() and 'VOLTAGE' in cell_v.upper():
                if "CHANNEL2" in cell_v.upper() and "VOLTAGE" in cell_v.upper():
                    avdd_col = col.column
                if "MODE" in cell_v.upper():
                    mode_col = col.column
                if "REGISTER_1" in cell_v.upper():
                    vref_col = col.column
                if "TEST LOG" in cell_v.upper():
                    log_col = col.column
        # last_c = sheet.max_column
        last_c = log_col
        if not log_col:
            print(f'"Test Log" not find in the sheet{ws}')
        else:
            sub_1 = f"EYE_FIG_REAL" if detail_m == 1 else "EYE_FIG"
            if demo_mode != 1:
                try:
                    path = os.path.join(path, sub_1)
                    os.makedirs(path)
                except Exception as e:
                    print(e)
                    pass
            for row in range(ws.max_row + 1 - start_row):
                eye_sn = [
                    [[[0] * 1 for _ in range(8)] for _ in range(4)] for _ in range(4)
                ]
                chk_row = ws.max_row - row if demo_mode == 1 else start_row + row
                if demo_mode == 1:
                    if not sheet.cell(row=chk_row, column=chip_col).value:
                        continue
                if sel_chip != "":
                    if sheet.cell(row=chk_row, column=chip_col).value != sel_chip:
                        continue
                if sel_mode != "":
                    if sheet.cell(row=chk_row, column=mode_col).value != sel_mode:
                        continue
                if sel_temper != "":
                    if (
                        str(sheet.cell(row=chk_row, column=temper_col).value)
                        != sel_temper
                    ):
                        continue
                if sel_avdd != "":
                    if not sheet.cell(row=chk_row, column=avdd_col).value:
                        continue
                    if float(sel_avdd) != float(
                        sheet.cell(row=chk_row, column=avdd_col).value
                    ):
                        continue
                chip_get = sheet.cell(row=chk_row, column=chip_col).value
                mode_get = sheet.cell(row=chk_row, column=mode_col).value
                temper_get = sheet.cell(row=chk_row, column=temper_col).value
                avdd_get = sheet.cell(row=chk_row, column=avdd_col).value
                path_val = ws.cell(row=chk_row, column=log_col).value
                if path_val != None:
                    txt_path = path_val.split('"')[1]
                    print(f"row={chk_row} path={txt_path}", flush=True)
                    Full_Eye_Results_list = self.txt_log_full_eye_check(
                        txt_path=txt_path, mode=detail_m, off_aply=off_aply
                    )
                else:
                    print(f"row={chk_row} no path find", flush=True)

                # "win_left Excel's col num"_"Die and slice name"
                path_name = f"ROW{chk_row}_{chip_get}_{mode_get}_T{temper_get}_{int(float(avdd_get) * 1000)}mV"
                rel_sp = path if demo_mode == 1 else f".\\{sub_1}\\{path_name}\\"
                sub_path = path if demo_mode == 1 else os.path.join(path, path_name)
                if demo_mode != 1:
                    try:
                        os.makedirs(sub_path)
                    except Exception as e:
                        print(e)
                        pass
                sp = sub_path if demo_mode == 1 else sub_path + "\\"
                # HYPERLINK = "=HYPERLINK(\"" + sp + "\"," + "\"" + "Fig_Folder" + "\"" + ")"
                HYPERLINK = (
                    '=HYPERLINK("' + rel_sp + '",' + '"' + "Fig_Folder" + '"' + ")"
                )
                sheet.cell(row=chk_row, column=last_c + 1).value = HYPERLINK
                sheet.cell(row=chk_row, column=last_c + 1).font = font
                sheet.cell(row=chk_row, column=last_c + 1).alignment = align
                sheet.cell(row=chk_row, column=last_c + 1).border = border
                cn = 1
                die_ck = []
                die_sn = -1
                for one_eye_data in Full_Eye_Results_list:
                    die_n = one_eye_data[0].split(" G")[0]
                    group_n = one_eye_data[0].split(" G")[1]
                    group_n = group_n.split(" S#")[0]
                    slice_n = one_eye_data[0].split(" S#")[1]
                    slice_n = slice_n.split(" O")[0]
                    offset_n = one_eye_data[0].split(" O")[1]
                    eye_sn[int(die_n)][int(group_n)][int(slice_n)][0] += 1
                    eye_sn_v = eye_sn[int(die_n)][int(group_n)][int(slice_n)][0]
                    end_txt = "" if demo_mode == 1 else f"{offset_n}_{eye_sn_v}"
                    Title = f"D{die_n}G{group_n}S{slice_n}{end_txt}"
                    f_Title = f"{chip_get}_{mode_get}_T{temper_get}_{int(float(avdd_get) * 1000)}mV_{Title}"
                    slice_num = slice_n
                    self.Graph_shmoo_vref_plot(
                        train_mode=1 + detail_m,
                        sheet=sheet,
                        sel_avdd=float(avdd_get),
                        vref_list=range(0, 64, 2),
                        slice_num=int(slice_num),
                        path=sp,
                        Title=Title,
                        f_Title=f_Title,
                        train_left=one_eye_data[1],
                        train_right=one_eye_data[2],
                        eye_detail=one_eye_data[4],
                    )
                    cn += 1
                    if die_n not in die_ck:
                        die_ck.append(die_n)
                        die_sn += 1
                        die_w = die_sn
                    else:
                        die_w = die_ck.index(die_n)
                    c_idx = die_w * 8 + int(slice_n) + 2
                    HYPERLINK = (
                        '=HYPERLINK("'
                        + rel_sp
                        + "Vref Window_"
                        + f_Title
                        + ".png"
                        + '",'
                        + '"'
                        + f"{eye_sn_v}_{one_eye_data[3][0]}/{one_eye_data[3][1]}_"
                        f"{one_eye_data[3][2]}/{one_eye_data[3][3]}" + '"' + ")"
                    )
                    sheet.cell(row=chk_row, column=last_c + c_idx).value = HYPERLINK
                    sheet.cell(row=chk_row, column=last_c + c_idx).font = font
                    sheet.cell(row=chk_row, column=last_c + c_idx).alignment = align
                    sheet.cell(row=chk_row, column=last_c + c_idx).border = border
                self.xls_save(work_book, path=xls_path)
                if demo_mode == 1 and Full_Eye_Results_list:
                    break
            print(f"\034Graph generated end for the sheet{ws}")
        work_book.close()

        # for i in range(len(Test_Result)):
        #     if Test_Result[i] == 'Hyperlink_Log':
        #         font = Font(u'Arial', size=10, bold=False, italic=False, strike=False, color='000000')
        #         HYPERLINK = "=HYPERLINK(\"" + Hyperlink_path + "\"," + "\"" + Test_Result[i] + "\"" + ")"
        #         Select_sheet.cell(row=Start_row_Num, column=i + 1).value = HYPERLINK
        #         Select_sheet.cell(row=Start_row_Num, column=i + 1).font = font
        #     else:
        #         font = Font(u'Arial', size=10, bold=False, italic=False, strike=False, color='000000')
        #         Select_sheet.cell(row=Start_row_Num, column=i + 1).font = font
        #         Select_sheet.cell(row=Start_row_Num, column=i + 1).value = Test_Result[i]
        #     Select_sheet.cell(row=Start_row_Num, column=i + 1).alignment = align
        #     Select_sheet.cell(row=Start_row_Num, column=i + 1).border = border
        # # print(Test_Result)
        # Excel_Path.save(excel_path)

    def Graph_shmoo_vref_BK2(self, **kargs):
        path = kargs.get("path", self.gui.shmoo_load_path.GetPath())
        # NV_voltage = kargs.get('NV_voltage', 0.75)
        # sheet = kargs.get('sheet_type', '')

        # type = kargs.get('type', 0.75)

        work_book = openpyxl.load_workbook(path)
        fp_name = path.split("\\")
        fname = fp_name[-1]
        path = path.rstrip(fname)
        sheet = work_book[self.gui.shmoo_type.Value]
        sel_chip = self.gui.shmoo_chip_number.Value
        sel_mode = self.gui.shmoo_chip_mode.Value
        sel_temper = self.gui.shmoo_test_temperature.Value
        sel_avdd = self.gui.shmoo_VDDC_voltage.Value

        chip_col = None
        temper_col = None
        avdd_col = None
        mode_col = None
        vref_col = None
        win_left_col = []
        title_list = []
        for col in sheet["4:4"]:
            cell_v = sheet.cell(row=col.row, column=col.column).value
            if cell_v is not None:
                title_list.append(cell_v)
                if "TRAIN" in cell_v.upper() and "WIN_LEFT" in cell_v.upper():
                    train_win_col = col.column
                    win_left_col.append(train_win_col)
        for col in sheet["5:5"]:
            cell_v = sheet.cell(row=col.row, column=col.column).value
            if cell_v is not None:
                title_list.append(cell_v)
                if "CHIP NUM" in cell_v.upper():
                    chip_col = col.column
                if "TEMPERATURE" in cell_v.upper():
                    temper_col = col.column
                if "CHANNEL1" in cell_v.upper() and "VOLTAGE" in cell_v.upper():
                    avdd_col = col.column
                if "MODE" in cell_v.upper():
                    mode_col = col.column
                if "REGISTER_1" in cell_v.upper():
                    vref_col = col.column

        vref_list = []
        check_row = []
        for z in range(6, sheet.max_row + 1):
            if (
                sel_chip == sheet.cell(row=z, column=chip_col).value
                and sel_mode == sheet.cell(row=z, column=mode_col).value
                and sel_temper == str(sheet.cell(row=z, column=temper_col).value)
            ):
                if float(sel_avdd) == sheet.cell(row=z, column=avdd_col).value:
                    vref_get = sheet.cell(row=z, column=vref_col).value
                    if "0X" in vref_get.upper():
                        base = 16
                    else:
                        base = 10
                    vref_list.append(int(vref_get, base))
                    check_row.append(z)

        if len(check_row) == 0:
            print(f"no data for plot")
            return
        # "win_left Excel's col num"_"Die and slice name"
        for slice_sel in [
            "Die1-Slice-0",
            "Die1-Slice-1",
            "Die1-Slice-2",
            "Die1-Slice-3",
            "Die1-Slice-4",
            "Die1-Slice-5",
            "Die1-Slice-6",
            "Die1-Slice-7",
            "Die0-Slice-0",
            "Die0-Slice-1",
            "Die0-Slice-2",
            "Die0-Slice-3",
            "Die0-Slice-4",
            "Die0-Slice-5",
            "Die0-Slice-6",
            "Die0-Slice-7",
        ]:
            slice_n = int((slice_sel.split("-"))[2])
            if "Die1" in slice_sel:
                die_chk = 0
            else:
                die_chk = 1
            f_Title = f"{sel_chip}_{sel_mode}_T{sel_temper}_{int(float(sel_avdd) * 1000)}mV_{slice_sel}"
            Title = f"{slice_sel}"

            fig, gnt = plt.subplots()  # Declaring a figure "gnt"
            gnt.set_ylim(5, 45)  # Setting Y-axis limits
            gnt.set_xlim(-32 / 64, 32 / 64)  # Setting X-axis limits

            # Setting labels for x-axis and y-axis
            # font = {'family': 'Arial', 'weight': 'normal', 'size': 16}
            font = {"family": "Arial", "weight": "normal", "size": 4}
            plt.title("Vref Window_" + Title, font)
            gnt.set_xlabel("UI (Unit Interval) ", font)
            gnt.set_ylabel("Vref Voltage Value (mV)", font)
            # plt.yticks(fontsize=6)
            # plt.xticks(fontsize=10)
            plt.yticks(fontsize=2.5)
            plt.xticks(fontsize=2.5)

            slice_num = win_left_col[die_chk] + slice_n

            yticks = []
            for y in range(len(check_row)):
                Num = 15 + (10 * y)
                yticks += [Num]
            plt.yticks(yticks)
            x = [-0.5, -0.4, -0.3, -0.2, -0.1, 0, 0.1, 0.2, 0.3, 0.4, 0.5]
            plt.xticks(x)

            lable = []
            index = 0
            for reg_val in vref_list:
                if index % 2 == 0:
                    try:
                        voltage = reg_val * (sel_avdd / 64)
                        # voltage = round(voltage, 3)
                        voltage = int(voltage * 1000)
                        lable += [voltage]
                    except:
                        lable += [""]
                else:
                    lable += [""]
                index += 1
            gnt.set_yticklabels(lable)  # Labelling tickes of y-axis
            i_idx = 0
            for i in check_row:
                i_idx = i_idx + 1  # -6 to get correct idx (5 rows for title)
                Win_L = (
                    int(sheet.cell(row=i, column=slice_num).value) - 32
                ) / 64  # -31 is set_bar_center lable need 0
                Win_R = (
                    int(sheet.cell(row=i, column=slice_num + 8).value) - 32
                ) / 64  # -31 is set_bar_center lable need 0
                if Win_L >= Win_R:
                    gnt.broken_barh([(-1, 99)], (10 * i_idx, 10), facecolors="tab:red")
                else:
                    gnt.broken_barh(
                        [(-32 / 64, (abs(-32 / 64 - Win_L)))],
                        (10 * i_idx, 10),
                        facecolors="tab:red",
                    )
                    gnt.broken_barh(
                        [(Win_L, (abs(Win_R - Win_L)))],
                        (10 * i_idx, 10),
                        facecolors="tab:green",
                    )
                    gnt.broken_barh(
                        [(Win_R, 99)], (10 * i_idx, 10), facecolors="tab:red"
                    )
                if Win_L == None:
                    break

            figure = plt.gcf()
            # figure.set_size_inches(2.8, 2.1)
            figure.set_size_inches(3.3, 2.1)
            # figure.set_size_inches(22, 14)
            plt.savefig(path + "Vref Window_" + f_Title + ".png", dpi=200)
            print("Vref Window_" + f_Title)
            # plt.show()
            plt.close()

    def Graph_shmoo_vref_BK(self, **kargs):
        path = kargs.get("path", 0)
        NV_voltage = kargs.get("NV_voltage", 0.75)
        Sheet_Name = kargs.get("Sheet_Name", "Shmoo_vref")

        # "win_left Excel's col num"_"Die and slice name"
        for slice in [
            "19_Die0-Slice-1",
            "20_Die0-Slice-2",
            "21_Die0-Slice-3",
            "22_Die0-Slice-4",
            "23_Die0-Slice-5",
            "24_Die0-Slice-6",
            "25_Die0-Slice-7",
            "26_Die0-Slice-8",
            "35_Die1-Slice-1",
            "36_Die1-Slice-2",
            "37_Die1-Slice-3",
            "38_Die1-Slice-4",
            "39_Die1-Slice-5",
            "40_Die1-Slice-6",
            "41_Die1-Slice-7",
            "42_Die1-Slice-8",
        ]:
            slice_num = int((slice.split("_"))[0])
            Title = str((slice.split("_"))[1])

            fig, gnt = plt.subplots()  # Declaring a figure "gnt"
            gnt.set_ylim(5, 45)  # Setting Y-axis limits
            gnt.set_xlim(-32 / 64, 32 / 64)  # Setting X-axis limits

            # Setting labels for x-axis and y-axis
            font = {"family": "Arial", "weight": "normal", "size": 16}
            plt.title("Vref Window_" + Title, font)
            gnt.set_xlabel("UI (Unit Interval) ", font)
            gnt.set_ylabel("Vref Voltage Value", font)

            work_book = openpyxl.load_workbook(path)
            sheet = work_book[Sheet_Name]

            plt.yticks(fontsize=6)
            plt.xticks(fontsize=10)

            Reg_Num = sheet.max_row - 5  # 5 : excel row1 to row5(title)
            Num = yticks = []
            for y in range(Reg_Num):
                Num = 15 + (10 * y)
                yticks += [Num]
            plt.yticks(yticks)
            x = [-0.5, -0.4, -0.3, -0.2, -0.1, 0, 0.1, 0.2, 0.3, 0.4, 0.5]
            plt.xticks(x)

            lable = Win_L = []
            for z in range(Reg_Num):
                reg_val = sheet.cell(row=z + 6, column=15).value
                try:
                    voltage = (int(reg_val, 16)) * (NV_voltage / 64)
                    voltage = round(voltage, 3)
                    lable += [voltage]
                except:
                    pass
            gnt.set_yticklabels(lable)  # Labelling tickes of y-axis

            for i in range(Reg_Num):
                Win_L = (
                    int(sheet.cell(row=i + 6, column=slice_num).value) - 32
                ) / 64  # -31 is set_bar_center lable need 0
                Win_R = (
                    int(sheet.cell(row=i + 6, column=slice_num + 8).value) - 32
                ) / 64  # -31 is set_bar_center lable need 0
                if Win_L >= Win_R:
                    gnt.broken_barh(
                        [(-1, 99)], (10 * (i + 1), 10), facecolors="tab:red"
                    )
                else:
                    gnt.broken_barh(
                        [(-32 / 64, (abs(-32 / 64 - (Win_L))))],
                        (10 * (i + 1), 10),
                        facecolors="tab:red",
                    )
                    gnt.broken_barh(
                        [(Win_L, (abs(Win_R - Win_L)))],
                        (10 * (i + 1), 10),
                        facecolors="tab:green",
                    )
                    gnt.broken_barh(
                        [(Win_R, 99)], (10 * (i + 1), 10), facecolors="tab:red"
                    )
                if Win_L == None:
                    break

            figure = plt.gcf()

            figure.set_size_inches(22, 14)
            plt.savefig("Test Graph(EY0012A)/Vref Window_" + Title + ".png", dpi=200)
            print("Vref Window_" + Title)
            # plt.show()
            plt.close()

    def Graph_shmoo_vref_sheet(self, **kargs):  # Manual sheet use
        path = kargs.get("path", self.gui.shmoo_load_path.GetPath())
        # NV_voltage = kargs.get('NV_voltage', 0.75)
        # sheet = kargs.get('sheet_type', '')

        # type = kargs.get('type', 0.75)

        work_book = openpyxl.load_workbook(path)
        fp_name = path.split("\\")
        fname = fp_name[-1]
        path = path.rstrip(fname)
        sheet = work_book[self.gui.shmoo_type.Value]
        sel_chip = self.gui.shmoo_chip_number.Value
        sel_mode = self.gui.shmoo_chip_mode.Value
        sel_temper = self.gui.shmoo_test_temperature.Value
        sel_avdd = self.gui.shmoo_VDDC_voltage.Value

        chip_col = None
        temper_col = None
        avdd_col = None
        mode_col = None
        vref_col = None
        win_left_col = []
        title_list = []
        for col in sheet["4:4"]:
            cell_v = sheet.cell(row=col.row, column=col.column).value
            if cell_v is not None:
                title_list.append(cell_v)
                if "TRAIN" in cell_v.upper() and "WIN_LEFT" in cell_v.upper():
                    train_win_col = col.column
                    win_left_col.append(train_win_col)
        for col in sheet["5:5"]:
            cell_v = sheet.cell(row=col.row, column=col.column).value
            if cell_v is not None and type(cell_v) == str:
                title_list.append(cell_v)
                if "CHIP NUM" in cell_v.upper():
                    chip_col = col.column
                if "TEMPERATURE" in cell_v.upper():
                    temper_col = col.column
                if "CHANNEL1" in cell_v.upper() and "VOLTAGE" in cell_v.upper():
                    avdd_col = col.column
                if "MODE" in cell_v.upper():
                    mode_col = col.column
                if "REGISTER_2" in cell_v.upper():
                    vref_col = col.column

        vref_list = []
        check_row = []
        for z in range(6, sheet.max_row + 1):
            if (
                sel_chip == sheet.cell(row=z, column=chip_col).value
                and sel_mode == sheet.cell(row=z, column=mode_col).value
                and sel_temper == str(sheet.cell(row=z, column=temper_col).value)
            ):
                if float(sel_avdd) == sheet.cell(row=z, column=avdd_col).value:
                    vref_item = (sheet.cell(row=z, column=vref_col).value).split(",")
                    for ckitem in vref_item:
                        if "vref" in ckitem:
                            vref_get = ckitem.split("=")[1]
                            if "0X" in vref_get.upper():
                                base = 16
                            else:
                                base = 10
                            vref_list.append(int(vref_get, base))
                            check_row.append(z)

        if len(check_row) == 0:
            print(f"no data for plot")
            return
        # "win_left Excel's col num"_"Die and slice name"
        for slice_sel in [
            "DieA-Slice-0",
            "DieA-Slice-1",
            "DieA-Slice-2",
            "DieA-Slice-3",
            "DieA-Slice-4",
            "DieA-Slice-5",
            "DieA-Slice-6",
            "DieA-Slice-7",
            "DieB-Slice-0",
            "DieB-Slice-1",
            "DieB-Slice-2",
            "DieB-Slice-3",
            "DieB-Slice-4",
            "DieB-Slice-5",
            "DieB-Slice-6",
            "DieB-Slice-7",
        ]:
            slice_n = int((slice_sel.split("-"))[2])
            if "DieA" in slice_sel:
                die_chk = 0
            else:
                die_chk = 1
            Title = f"{sel_chip}_{sel_mode}_T{sel_temper}_{int(float(sel_avdd) * 1000)}mV_{slice_sel}"
            slice_num = win_left_col[die_chk] + slice_n
            self.Graph_shmoo_vref_plot(
                check_row=check_row,
                sheet=sheet,
                sel_avdd=sel_avdd,
                vref_list=vref_list,
                slice_num=slice_num,
                path=path,
                Title=Title,
                f_Title=Title,
                ytick_mod=1,
            )
            # # figure = plt.gcf()
            # figure = plt.figure()
            # fig, gnt = plt.subplots()  # Declaring a figure "gnt"
            # gnt.set_ylim(5, 45)  # Setting Y-axis limits
            # gnt.set_xlim(-32 / 64, 32 / 64)  # Setting X-axis limits
            #
            # # Setting labels for x-axis and y-axis
            # font = {'family': 'Arial', 'weight': 'normal', 'size': 16}
            # plt.title('Vref Window_' + Title, font)
            # gnt.set_xlabel('UI (Unit Interval) ', font)
            # gnt.set_ylabel('Vref Voltage Value (mV)', font)
            # plt.yticks(fontsize=6)
            # plt.xticks(fontsize=10)
            #
            # yticks = []
            # for y in range(len(check_row)):
            #     Num = 15 + (10 * y)
            #     yticks += [Num]
            # plt.yticks(yticks)
            # x = [-0.5, -0.4, -0.3, -0.2, -0.1, 0, 0.1, 0.2, 0.3, 0.4, 0.5]
            # plt.xticks(x)
            #
            # lable = []
            # for reg_val in vref_list:
            #     try:
            #         voltage = reg_val * (sel_avdd / 64)
            #         # voltage = round(voltage, 3)
            #         voltage = int(voltage*1000)
            #         lable += [voltage]
            #     except:
            #         pass
            # gnt.set_yticklabels(lable)  # Labelling tickes of y-axis
            # i_idx = 0
            # for i in check_row:
            #     i_idx = i_idx + 1  # -6 to get correct idx (5 rows for title)
            #     Win_L = (int(sheet.cell(row=i, column=slice_num).value) - 32) / 64  # -31 is set_bar_center lable need 0
            #     Win_R = (int(sheet.cell(row=i, column=slice_num + 8).value) - 32) / 64  # -31 is set_bar_center lable need 0
            #     if Win_L >= Win_R:
            #         gnt.broken_barh([(-1, 99)], (10 * i_idx, 10), facecolors='tab:red')
            #     else:
            #         gnt.broken_barh([(-32 / 64, (abs(-32 / 64 - Win_L)))], (10 * i_idx, 10), facecolors='tab:red')
            #         gnt.broken_barh([(Win_L, (abs(Win_R - Win_L)))], (10 * i_idx, 10), facecolors='tab:green')
            #         gnt.broken_barh([(Win_R, 99)], (10 * i_idx, 10), facecolors='tab:red')
            #     if Win_L == None:
            #         break
            #
            #
            # figure.set_size_inches(22, 14)
            #
            # plt.savefig(path + 'Vref Window_' + Title + '.jpg', dpi=200)
            # print('Vref Window_' + Title)
            # # plt.show()
            # plt.close()

    def Graph_shmoo_vref_plot(self, **kargs):
        train_mode = kargs.get("train_mode", 0)
        train_left = kargs.get("train_left", [-1] * 32)
        train_right = kargs.get("train_right", [-1] * 32)
        eye_detail = kargs.get("eye_detail", [[0] * 32] * 32)
        check_row = kargs.get("check_row", [6])
        sheet = kargs.get("sheet", None)
        sel_avdd = kargs.get("sel_avdd", 0.75)
        vref_list = kargs.get("vref_list", [32])
        slice_num = kargs.get("slice_num", 0)
        f_Title = kargs.get("f_Title", "")
        Title = kargs.get("Title", "")
        path = kargs.get("path", self.gui.shmoo_load_path.GetPath())
        ytick_mod = kargs.get("ytick_mod", 2)

        fig, gnt = plt.subplots()  # Declaring a figure "gnt"
        gnt.set_ylim(5, 45)  # Setting Y-axis limits
        gnt.set_xlim(-32 / 64, 32 / 64)  # Setting X-axis limits

        # Setting labels for x-axis and y-axis
        # font = {'family': 'Arial', 'weight': 'normal', 'size': 16}
        font = {"family": "Arial", "weight": "normal", "size": 4}
        plt.title("Vref Window_" + Title, font)
        gnt.set_xlabel("UI (Unit Interval) ", font)
        gnt.set_ylabel("Vref Voltage Value (mV)", font)
        # plt.yticks(fontsize=6)
        # plt.xticks(fontsize=10)
        plt.yticks(fontsize=2.5)
        plt.xticks(fontsize=2.5)

        figure = plt.gcf()
        # figure.set_size_inches(2.8, 2.1)
        figure.set_size_inches(3.3, 2.1)
        # figure.set_size_inches(22, 14)

        # slice_num = win_left_col[die_chk] + slice_n

        if train_mode == 1 or train_mode == 2:
            check_row = [6] * 32
            vref_list = range(0, 64, 2)

        yticks = []
        for y in range(len(check_row)):
            Num = 15 + (10 * y)
            yticks += [Num]
        plt.yticks(yticks)
        x = [-0.5, -0.4, -0.3, -0.2, -0.1, 0, 0.1, 0.2, 0.3, 0.4, 0.5]
        plt.xticks(x)

        lable = []
        index = 0
        for reg_val in vref_list:
            if index % ytick_mod == 0:
                try:
                    voltage = reg_val * (sel_avdd / 64)
                    # voltage = round(voltage, 3)
                    voltage = int(voltage * 1000)
                    lable += [voltage]
                except:
                    lable += [""]
            else:
                lable += [""]
            index += 1
        gnt.set_yticklabels(lable)  # Labelling tickes of y-axis
        if train_mode == 1:
            i_idx = 0
            for i in range(32):
                i_idx = i_idx + 1
                Win_L = (train_left[i] - 31) / 64  # -31 is set_bar_center lable need 0
                Win_R = (train_right[i] - 31) / 64  # -31 is set_bar_center lable need 0
                if Win_L >= Win_R:
                    gnt.broken_barh([(-1, 99)], (10 * i_idx, 10), facecolors="tab:red")
                else:
                    gnt.broken_barh(
                        [(-32 / 64, (abs(-32 / 64 - Win_L)))],
                        (10 * i_idx, 10),
                        facecolors="tab:red",
                    )
                    gnt.broken_barh(
                        [(Win_L, (abs(Win_R - Win_L)))],
                        (10 * i_idx, 10),
                        facecolors="tab:green",
                    )
                    gnt.broken_barh(
                        [(Win_R, 99)], (10 * i_idx, 10), facecolors="tab:red"
                    )
                if Win_L == None:
                    break
        elif train_mode == 2:
            i_idx = 0
            for i in range(32):
                i_idx = i_idx + 1
                xp = len(eye_detail[i]) - 1
                for k in range(xp):
                    # plot_f = 'tab:green' if eye_detail[i][k] == 1 else 'tab:red'
                    if eye_detail[i][k] == max(max(eye_detail)):
                        plot_f = "tab:green"
                    elif eye_detail[i][k] == 0:
                        plot_f = "tab:red"
                    else:
                        plot_f = "y"
                    gnt.broken_barh(
                        [((k - xp / 2) / xp, 1 / xp)],
                        (10 * i_idx, 10),
                        facecolors=plot_f,
                    )

        elif sheet is not None:
            i_idx = 0
            for i in check_row:
                i_idx = i_idx + 1  # -6 to get correct idx (5 rows for title)
                Win_L = (
                    int(sheet.cell(row=i, column=slice_num).value) - 31
                ) / 64  # -31 is set_bar_center lable need 0
                Win_R = (
                    int(sheet.cell(row=i, column=slice_num + 8).value) - 31
                ) / 64  # -31 is set_bar_center lable need 0
                if Win_L >= Win_R:
                    gnt.broken_barh([(-1, 99)], (10 * i_idx, 10), facecolors="tab:red")
                else:
                    gnt.broken_barh(
                        [(-32 / 64, (abs(-32 / 64 - Win_L)))],
                        (10 * i_idx, 10),
                        facecolors="tab:red",
                    )
                    gnt.broken_barh(
                        [(Win_L, (abs(Win_R - Win_L)))],
                        (10 * i_idx, 10),
                        facecolors="tab:green",
                    )
                    gnt.broken_barh(
                        [(Win_R, 99)], (10 * i_idx, 10), facecolors="tab:red"
                    )
                if Win_L == None:
                    break

        plt.savefig(path + "Vref Window_" + f_Title + ".png", dpi=200)
        print("Vref Window_" + f_Title)
        # plt.show()
        plt.close()
